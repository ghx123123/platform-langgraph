import csv
import hashlib
import html
import io
import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Callable, Iterable
from uuid import UUID, uuid4, uuid5

from backend.course_archives.models import ArchiveManifestItem, PrepareArchiveRequest
from backend.documents.service import parse_document
from backend.documents.storage import original_path, persist_original


PARSEABLE_EXTENSIONS = {
    ".pdf", ".docx", ".pptx", ".md", ".txt", ".xlsx", ".xls",
    ".csv", ".html", ".htm", ".py", ".json", ".ipynb",
}
PREVIEW_EXTENSIONS = {".pdf", ".docx", ".pptx"}

_CATEGORY_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("syllabus", ("教学大纲", "课程大纲", "课程标准", "syllabus")),
    ("schedule", ("进度表", "教学进度", "授课计划", "教学日历", "schedule")),
    ("teaching_record", ("教学记录", "迭代记录", "备课记录", "第1轮", "第2轮", "第3轮", "第4轮", "提示词")),
    ("review", ("审核", "审查", "评审", "督导", "质量报告", "辩论", "audit", "review")),
    ("experiment", ("实验", "实训", "实验指导", "experiment")),
    ("lesson_plan", ("教案", "教学设计", "讲稿", "说课", "lesson plan")),
    ("textbook", ("教材", "课本", "教科书", "textbook")),
    ("courseware", ("课件", "ppt", "powerpoint", "幻灯片")),
    ("interactive", ("仿真", "模拟", "交互", "可视化", "动画", "simulation")),
    ("reference", ("参考", "文献", "论文", "案例库", "reference")),
)

_MEDIA_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg",
    ".mp4", ".avi", ".mov", ".mkv", ".mp3", ".wav", ".flac",
}
_CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c", ".cpp",
    ".h", ".cs", ".m", ".ino", ".ipynb", ".json", ".yaml", ".yml",
}
_CHINESE_NUMBERS = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def normalize_archive_path(value: str) -> str:
    value = value.replace("\\", "/").strip().lstrip("/")
    parts = [part for part in PurePosixPath(value).parts if part not in {"", ".", ".."}]
    return "/".join(parts)[:500]


def classify_material(path: str) -> str:
    normalized = normalize_archive_path(path)
    lowered = normalized.lower()
    extension = Path(normalized).suffix.lower()
    for category, markers in _CATEGORY_RULES:
        if any(marker.lower() in lowered for marker in markers):
            return category
    if extension in _MEDIA_EXTENSIONS:
        return "media"
    if extension in _CODE_EXTENSIONS:
        return "code"
    if extension in {".ppt", ".pptx"}:
        return "courseware"
    if extension in {".html", ".htm"}:
        return "interactive"
    if extension == ".pdf" and re.search(r"第\s*(?:[0-9]{1,2}|[一二三四五六七八九十])\s*章", normalized):
        return "textbook"
    if extension == ".docx" and "ocr" in lowered and re.search(r"第\s*(?:[0-9]{1,2}|[一二三四五六七八九十])\s*章", normalized):
        return "textbook"
    return "other"


def extract_chapter(path: str) -> str | None:
    normalized = normalize_archive_path(path)
    match = re.search(r"第\s*([0-9]{1,2}|[一二三四五六七八九十])\s*章", normalized)
    if match:
        token = match.group(1)
        number = int(token) if token.isdigit() else _CHINESE_NUMBERS.get(token)
        return f"第{number}章" if number else f"第{token}章"
    match = re.search(r"(?:chapter|chap|ch)[\s_-]*([0-9]{1,2})", normalized, re.IGNORECASE)
    if match:
        return f"第{int(match.group(1))}章"
    return None


def extract_lesson(path: str) -> str | None:
    normalized = normalize_archive_path(path)
    match = re.search(r"第\s*([0-9]{1,2})\s*(讲|次课|课)", normalized)
    if match:
        return f"第{int(match.group(1))}{match.group(2)}"
    match = re.search(r"([0-9]{1,2})\s*[-~至]\s*([0-9]{1,2})\s*课时", normalized)
    if match:
        return f"{match.group(1)}-{match.group(2)}课时"
    return None


def extract_version(path: str) -> str | None:
    name = Path(normalize_archive_path(path)).stem
    markers: list[str] = []
    round_match = re.search(r"第\s*([0-9]{1,2})\s*轮", name)
    if round_match:
        markers.append(f"第{int(round_match.group(1))}轮")
    for label, pattern in (
        ("定稿", r"最终|定稿|终稿|提交版"),
        ("优化版", r"优化|改进|新版"),
        ("OCR版", r"ocr|文字识别"),
        ("副本", r"副本|复制|copy|\([0-9]+\)$"),
    ):
        if re.search(pattern, name, re.IGNORECASE):
            markers.append(label)
    return " / ".join(dict.fromkeys(markers)) or None


def _plain_text(filename: str, data: bytes) -> tuple[str, str]:
    extension = Path(filename).suffix.lower()
    if extension in {".xlsx", ".xls"}:
        text, _ = _spreadsheet_content(extension, data)
        return text, "spreadsheet"
    decoded = data.decode("utf-8-sig", errors="replace")
    if extension in {".html", ".htm"}:
        decoded = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", decoded, flags=re.IGNORECASE | re.DOTALL)
        decoded = re.sub(r"<[^>]+>", "\n", decoded)
        return html.unescape(decoded), "HTML text"
    if extension == ".ipynb":
        try:
            notebook = json.loads(decoded)
            cells = []
            for cell in notebook.get("cells", []):
                source = cell.get("source", [])
                cells.append("".join(source) if isinstance(source, list) else str(source))
            return "\n\n".join(cells), "Jupyter notebook"
        except json.JSONDecodeError:
            pass
    return decoded, "UTF-8 text"


def _spreadsheet_content(extension: str, data: bytes) -> tuple[str, dict]:
    lines: list[str] = []
    sheets: list[dict] = []
    if extension == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(data), read_only=False, data_only=True)
        for sheet in workbook.worksheets:
            lines.append(f"# 工作表：{sheet.title}")
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    lines.append("\t".join(values))
                    rows.append(values)
            sheets.append({
                "name": sheet.title,
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "merged_ranges": [str(item) for item in sheet.merged_cells.ranges],
            })
        workbook.close()
    else:
        import xlrd

        workbook = xlrd.open_workbook(file_contents=data)
        for sheet in workbook.sheets():
            lines.append(f"# 工作表：{sheet.name}")
            rows: list[list[str]] = []
            for row_index in range(sheet.nrows):
                values = [str(value).strip() for value in sheet.row_values(row_index)]
                if any(values):
                    lines.append("\t".join(values))
                    rows.append(values)
            merged_ranges = [f"R{rlo + 1}C{clo + 1}:R{rhi}C{chi}" for rlo, rhi, clo, chi in sheet.merged_cells]
            sheets.append({
                "name": sheet.name,
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "merged_ranges": merged_ranges,
            })
    return "\n".join(lines), {"kind": "spreadsheet", "sheets": sheets}


def _parse_uploaded(
    filename: str,
    data: bytes,
    document_store_path: Path,
    document_id: str | None = None,
    *,
    full_extraction: bool = False,
    engine: str = "mineru",
    progress_cb: Callable[[int, str], None] | None = None,
) -> dict:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        # Archive scans stay fast. A teacher explicitly importing a PDF into a
        # material unit gets the same OCR-capable parser used by course design.
        parsed = parse_document(filename, data, engine=engine, progress_cb=progress_cb) if full_extraction else _fast_pdf_document(filename, data)
    elif extension == ".docx":
        parsed = _fast_docx_document(filename, data)
    elif extension in {".xlsx", ".xls"}:
        text, structured = _spreadsheet_content(extension, data)
        parsed = parse_document(f"{Path(filename).stem}.txt", text.encode("utf-8"))
        parsed["file_name"] = Path(filename).name
        parsed["course_name"] = Path(filename).stem
        parsed["structured_content"] = structured
        parsed["extraction_report"]["format"] = extension.lstrip(".").upper()
        parsed["extraction_report"]["engine"] = "spreadsheet structure index"
        parsed["extraction_report"]["table_count"] = len(structured["sheets"])
    elif extension in {".pptx", ".md", ".txt"}:
        parsed = parse_document(Path(filename).name, data)
    else:
        text, engine = _plain_text(filename, data)
        parsed = parse_document(f"{Path(filename).stem}.txt", text.encode("utf-8"))
        parsed["file_name"] = Path(filename).name
        parsed["course_name"] = Path(filename).stem
        parsed["extraction_report"]["format"] = extension.lstrip(".").upper()
        parsed["extraction_report"]["engine"] = engine
    if document_id:
        parsed["document_id"] = document_id
    persist_original(document_store_path, parsed["document_id"], Path(filename).name, data)
    return parsed


def _fast_docx_document(filename: str, data: bytes) -> dict:
    """Index accessible Word text and tables without OCRing embedded images."""
    from docx import Document

    document = Document(io.BytesIO(data))
    blocks: list[str] = []
    paragraphs: list[dict] = []
    for paragraph in document.paragraphs:
        value = _clean_docx_line(paragraph.text)
        if value:
            paragraphs.append({"text": value, "style": paragraph.style.name if paragraph.style else ""})
            prefix = "# " if paragraph.style and paragraph.style.name.startswith("Title") else ""
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                level_match = re.search(r"\d+", paragraph.style.name)
                prefix = "#" * min(int(level_match.group()) if level_match else 2, 6) + " "
            elif re.match(r"^[一二三四五六七八九十]+[、,，.．]\s*[^。；，]{2,40}$", value):
                value = re.sub(r"^[一二三四五六七八九十]+[、,，.．]\s*", "", value)
                prefix = "## "
            elif re.match(r"^[0-9]{1,2}[.、]\s*[^。；，]{2,40}$", value):
                value = re.sub(r"^[0-9]{1,2}[.、]\s*", "", value)
                prefix = "### "
            blocks.append(prefix + value)
    tables: list[dict] = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            values = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            if any(values):
                rows.append("\t".join(values))
        if rows:
            blocks.append("\n".join(rows))
            tables.append({
                "index": table_index,
                "rows": [row.split("\t") for row in rows],
                "row_count": len(rows),
                "column_count": max((len(row.split("\t")) for row in rows), default=0),
            })
    chapter = extract_chapter(filename)
    first_plain = next((re.sub(r"^#+\s*", "", item) for item in blocks if item.strip()), "")
    if chapter and first_plain and not first_plain.startswith(chapter):
        first_title = first_plain if len(first_plain) <= 32 else Path(filename).stem
        first_title = re.sub(r"机电传动控制\s*第?2版|第\s*\d+\s*章|\(?OCR\)?", "", first_title, flags=re.IGNORECASE).strip(" _-()（）")
        blocks.insert(0, f"# {chapter}{' ' + first_title if first_title else ''}")
    text = "\n\n".join(blocks)
    if len(text.strip()) < 10:
        raise ValueError("Word 文档没有可访问正文；已保留清单，请作为单份材料导入后执行图片 OCR")
    parsed = parse_document(f"{Path(filename).stem}.txt", text.encode("utf-8"))
    parsed["file_name"] = Path(filename).name
    parsed["course_name"] = re.sub(r"^(第\s*[0-9一二三四五六七八九十]+\s*章)(?=\S)", r"\1 ", parsed["course_name"])
    parsed["structured_content"] = {"kind": "word", "paragraphs": paragraphs, "tables": tables}
    parsed["extraction_report"].update({
        "format": "DOCX",
        "engine": "python-docx 快速索引（无图片 OCR）",
        "table_count": len(document.tables),
        "image_count": len(document.inline_shapes),
    })
    if document.inline_shapes:
        parsed["extraction_report"]["warnings"] = [
            *parsed["extraction_report"].get("warnings", []),
            f"资料库快速索引未识别 {len(document.inline_shapes)} 张嵌入图片；选为主材料后可通过单文件导入执行图片 OCR",
        ]
    return parsed


def _clean_docx_line(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip()
    if not value:
        return ""
    # OCR books often repeat chapter titles with dotted page leaders on every page.
    if re.match(r"^第[一二三四五六七八九十0-9]+章\s+.+?[.．…·]+\s*(?:\d\s*)+$", value):
        return ""
    if re.match(r"^\d+\s*(?:[.．…·■]\s*)*机电传动控制(?:\s+第\s*\d+\s*版)?\s*$", value):
        return ""
    value = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", value)
    return value


def _fast_pdf_document(filename: str, data: bytes) -> dict:
    """Build the archive index from native PDF text without blocking on book-wide OCR.

    Full OCR remains available through the existing single-document import once the
    teacher chooses a scanned PDF as the preparation pack's primary material.
    """
    import fitz

    document = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    scanned_pages = 0
    try:
        for page_number, page in enumerate(document, start=1):
            text = page.get_text("text", sort=True).strip()
            if len(text) < 24:
                scanned_pages += 1
            pages.append(f"[第 {page_number} 页]\n{text}" if text else f"[第 {page_number} 页]\n（该页未提取到原生文字）")
    finally:
        document.close()
    combined = "\n\n".join(pages)
    if len(re.sub(r"\[第 \d+ 页\]|（该页未提取到原生文字）", "", combined).strip()) < 10:
        raise ValueError("扫描 PDF 未包含足够的原生文字；已保留清单和原文件，请作为单份材料导入后执行完整 OCR")
    parsed = parse_document(f"{Path(filename).stem}.txt", combined.encode("utf-8"))
    parsed["file_name"] = Path(filename).name
    report = parsed["extraction_report"]
    report.update({
        "format": "PDF",
        "engine": "PyMuPDF 快速索引（无 OCR）",
        "page_count": len(pages),
        "text_block_count": len(pages) - scanned_pages,
        "scanned_page_count": scanned_pages,
        "ocr_page_count": 0,
    })
    if scanned_pages:
        report["warnings"] = [
            *report.get("warnings", []),
            f"资料库快速索引发现 {scanned_pages} 页缺少原生文字；选为主材料后可通过单文件导入执行完整 OCR",
        ]
        report["quality_score"] = min(report.get("quality_score", 100), 68 if scanned_pages == len(pages) else 78)
        report["quality_level"] = "medium"
    return parsed


def _probable_duplicate_key(item: dict) -> str | None:
    if item["size"] <= 0:
        return None
    stem = Path(item["name"]).stem.lower()
    stem = re.sub(r"(?:副本|复制|copy|\([0-9]+\)|[_-]?[0-9]+)$", "", stem, flags=re.IGNORECASE)
    stem = re.sub(r"\s+", "", stem)
    return f"{stem}:{item['size']}" if stem else None


def _assign_duplicates(materials: list[dict]) -> int:
    exact: dict[str, list[dict]] = defaultdict(list)
    probable: dict[str, list[dict]] = defaultdict(list)
    for item in materials:
        if item.get("sha256"):
            exact[item["sha256"]].append(item)
        key = _probable_duplicate_key(item)
        if key:
            probable[key].append(item)
    groups = 0
    assigned: set[str] = set()
    for digest, items in exact.items():
        if len(items) < 2:
            continue
        groups += 1
        group = f"sha256:{digest[:12]}"
        for item in items:
            item["duplicate_group"] = group
            assigned.add(item["id"])
    for key, items in probable.items():
        remaining = [item for item in items if item["id"] not in assigned]
        if len(remaining) < 2:
            continue
        groups += 1
        group = f"probable:{hashlib.sha1(key.encode('utf-8')).hexdigest()[:12]}"
        for item in remaining:
            item["duplicate_group"] = group
    return groups


def _course_title(name: str, materials: list[dict], documents: dict[str, dict]) -> str:
    candidates: list[str] = []
    for item in materials:
        if item["category"] not in {"syllabus", "schedule", "lesson_plan"}:
            continue
        value = Path(item["name"]).stem
        value = re.split(r"教学大纲|课程大纲|进度表|教学进度|教案|教学设计", value, maxsplit=1)[0]
        value = re.sub(r"（[^）]*）|\([^)]*\)", "", value)
        value = re.sub(r"[-_+]?\d{2,}.*$", "", value).strip(" _-+")
        if 2 <= len(value) <= 40:
            candidates.append(value)
    if candidates:
        return Counter(candidates).most_common(1)[0][0]
    parsed_titles = [doc.get("course_name", "") for doc in documents.values()]
    parsed_titles = [title for title in parsed_titles if 2 <= len(title) <= 50 and not title.startswith("第")]
    return Counter(parsed_titles).most_common(1)[0][0] if parsed_titles else name


def _chapters(materials: list[dict]) -> list[dict]:
    grouped: dict[str, list[str]] = defaultdict(list)
    for item in materials:
        if item.get("chapter"):
            grouped[item["chapter"]].append(item["id"])
    def sort_key(value: str) -> int:
        match = re.search(r"\d+", value)
        return int(match.group()) if match else 999
    return [
        {"key": chapter, "label": chapter, "material_ids": ids, "material_count": len(ids)}
        for chapter, ids in sorted(grouped.items(), key=lambda pair: sort_key(pair[0]))
    ]


def _schedule_entries(materials: list[dict], documents: dict[str, dict]) -> list[dict]:
    entries: list[dict] = []
    for item in materials:
        if item["category"] != "schedule" or item["id"] not in documents:
            continue
        text = documents[item["id"]].get("raw_text", "")
        for line in text.splitlines():
            value = re.sub(r"\s+", " ", line).strip(" |\t")
            if not 8 <= len(value) <= 240:
                continue
            chapter = extract_chapter(value)
            lesson = re.search(r"(?:第\s*)?\d{1,2}(?:\s*[-~至]\s*\d{1,2})?\s*(?:周|次|讲|课时)", value)
            if not chapter and not lesson:
                continue
            label = lesson.group(0).replace(" ", "") if lesson else chapter or "教学安排"
            entries.append({
                "id": f"schedule-{len(entries) + 1}",
                "label": label,
                "content": value,
                "chapter": chapter,
                "source_material_id": item["id"],
            })
            if len(entries) >= 48:
                return entries
    return entries


def _habit(
    key: str,
    title: str,
    description: str,
    instruction: str,
    evidence: Iterable[str],
    confidence: float,
) -> dict | None:
    evidence_items = list(dict.fromkeys(evidence))[:4]
    if not evidence_items:
        return None
    return {
        "key": key,
        "title": title,
        "description": description,
        "reusable_instruction": instruction,
        "evidence": evidence_items,
        "confidence": confidence,
    }


def infer_preparation_habits(materials: list[dict], documents: dict[str, dict]) -> list[dict]:
    by_category: dict[str, list[dict]] = defaultdict(list)
    for item in materials:
        by_category[item["category"]].append(item)
    searchable: dict[str, str] = {}
    for item in materials:
        doc = documents.get(item["id"], {})
        searchable[item["id"]] = f"{item['path']}\n{doc.get('raw_text', '')[:16000]}".lower()
    habits: list[dict | None] = []
    alignment = by_category["syllabus"] + by_category["schedule"]
    habits.append(_habit(
        "constraint_alignment", "先定课程约束，再做讲次设计",
        "备课先以教学大纲、进度表和正式模板确定学时、目标与交付格式。",
        "先核对教学大纲、教学进度和正式模板；所有目标、内容范围、课时与成果格式必须与其一致。",
        (item["path"] for item in alignment), 0.96,
    ))
    round_groups: dict[str, list[dict]] = defaultdict(list)
    for item in by_category["teaching_record"]:
        if re.search(r"第\s*\d+\s*轮", item["path"]):
            round_groups[item.get("chapter") or str(PurePosixPath(item["path"]).parent)].append(item)
    repeated = [item for items in round_groups.values() if len(items) >= 2 for item in items]
    habits.append(_habit(
        "same_scope_iteration", "固定知识范围，多轮重复打磨",
        "连续轮次围绕同一章节或专题改进案例、讲法、活动与评价，而不是顺延新内容。",
        "每一轮都保持当前选定知识点不变；依据上一轮督导意见重做教学设计与讲授表达，不得自动进入下一章节。",
        (item["path"] for item in repeated), 0.98,
    ))
    engineering = [item for item in materials if re.search(r"工程|工业|案例|项目|机器人|电机|传感器", searchable[item["id"]])]
    habits.append(_habit(
        "engineering_cases", "用工程任务组织抽象知识",
        "优先使用真实设备、工业案例、项目任务和参数问题建立教学情境。",
        "用可验证的工程任务或工业案例引出知识点，明确问题链、参数、约束和预期产出，避免只罗列概念。",
        (item["path"] for item in engineering), 0.90,
    ))
    structure = [item for item in materials if re.search(r"对比|比较|知识脉络|逻辑链|系统链", searchable[item["id"]])]
    habits.append(_habit(
        "comparison_system_chain", "强调对比与系统逻辑链",
        "通过方案对比、知识脉络和系统输入输出关系帮助学生建立整体结构。",
        "教学设计中同时给出关键方案对比和系统逻辑链，说明概念之间的因果、输入输出与适用边界。",
        (item["path"] for item in structure), 0.88,
    ))
    obe = [item for item in materials if re.search(r"obe|成果导向|课程目标|达成度", searchable[item["id"]], re.IGNORECASE)]
    habits.append(_habit(
        "obe_mapping", "教学目标与 OBE 证据闭环",
        "目标、课堂任务和评价证据之间需要能够逐项对应。",
        "为每个教学目标配置可观察的课堂任务与评价证据，并说明其对应关系和达成标准。",
        (item["path"] for item in obe), 0.91,
    ))
    timed = [item for item in by_category["lesson_plan"] + by_category["schedule"] if re.search(r"分钟|学时|课时", searchable[item["id"]])]
    habits.append(_habit(
        "timeboxed_design", "按课时精确分配教学容量",
        "正式教案关注每个环节的分钟数、重点容量和讲练比例。",
        "按当前讲次总时长分配各教学环节，时间总和必须一致，并把更多时间留给选定难点和练习反馈。",
        (item["path"] for item in timed), 0.93,
    ))
    multi = by_category["courseware"] + by_category["experiment"] + by_category["code"] + by_category["interactive"]
    if len({item["category"] for item in multi}) >= 3:
        habits.append(_habit(
            "multi_format_outputs", "同一讲次组织多形态教学资料",
            "教案之外，还会配套课件、实验、代码和交互资源。",
            "围绕同一知识范围组织正式教案、课件要点、实验或代码资源和可视化素材，并标注各自的课堂用途。",
            (item["path"] for item in multi), 0.94,
        ))
    reviews = by_category["review"]
    habits.append(_habit(
        "quality_audit", "形成内容、公式与呈现质量审查",
        "每轮或章节完成后通过审核报告和督导意见识别高优先级问题。",
        "生成后按“优点、不足、建议”审查内容准确性、公式、案例适切性和版式，并给出下一轮可直接执行的优化提示词。",
        (item["path"] for item in reviews), 0.95,
    ))
    archive_docs = [item for item in by_category["lesson_plan"] if item["extension"] in {".doc", ".docx"}]
    habits.append(_habit(
        "formal_archival", "以正式 Word 教案归档",
        "多轮生成内容最终需要回收到规范教案或比赛支撑材料中。",
        "最终成果应按正式教案结构整理，保留目标、重难点、教学过程、评价、课后反思与改进记录，便于 Word 归档。",
        (item["path"] for item in archive_docs), 0.89,
    ))
    return [item for item in habits if item is not None]


def _profile_prompt(habits: list[dict]) -> str:
    lines = ["# 可复用备课经验", "以下规则来自本课程历史材料证据，后续每次教学设计均应遵守："]
    lines.extend(f"- {habit['title']}：{habit['reusable_instruction']}" for habit in habits)
    return "\n".join(lines)


def _summary(record: dict) -> dict:
    fields = (
        "id", "name", "course_title", "created_at", "updated_at", "total_files",
        "parsed_files", "duplicate_groups", "chapter_count", "categories",
    )
    return {
        **{field: record[field] for field in fields},
        "academic_term": record.get("academic_term") or "",
        "course_code": record.get("course_code") or "",
        "local_root": record.get("local_root"),
        "last_scanned_at": record.get("last_scanned_at"),
        "source_folder_count": len(record.get("source_folders", [])),
    }


def public_archive(record: dict) -> dict:
    return {key: value for key, value in record.items() if not key.startswith("_")}


def store_course_archive_originals(
    record: dict,
    uploads: list[tuple[str, bytes]],
    document_store_path: Path,
) -> tuple[dict, int, list[str]]:
    materials = {item.get("path", "").casefold(): item for item in record.get("materials", [])}
    documents = record.setdefault("_documents", {})
    replaced_document_ids: list[str] = []
    updated_count = 0
    for raw_path, data in uploads:
        path = normalize_archive_path(raw_path)
        material = materials.get(path.casefold())
        if not material:
            raise KeyError(f"目录清单中没有文件：{path}")
        digest = hashlib.sha256(data).hexdigest()
        if material.get("sha256") == digest and material.get("document_id"):
            continue
        previous_document_id = material.get("document_id")
        document_id = str(uuid4())
        persist_original(document_store_path, document_id, material["name"], data)
        if previous_document_id:
            replaced_document_ids.append(previous_document_id)
        documents.pop(material["id"], None)
        extension = material.get("extension", "")
        material.update({
            "sha256": digest,
            "document_id": document_id,
            "preview_available": extension in PREVIEW_EXTENSIONS,
            "parse_status": "metadata_only" if extension in PARSEABLE_EXTENSIONS else "unsupported",
            "parse_message": "原文件已重新加载；进入备课时再按需提取正文" if previous_document_id else "原文件已保存；进入备课时再按需提取正文",
            "character_count": 0,
            "excerpt": "",
        })
        updated_count += 1
    record["parsed_files"] = len(documents)
    record["updated_at"] = utc_now()
    return record, updated_count, replaced_document_ids


def analyze_course_archive(
    archive_name: str,
    manifest: list[ArchiveManifestItem],
    uploads: list[tuple[str, bytes]],
    document_store_path: Path,
    existing_record: dict | None = None,
    metadata: dict | None = None,
    extract_uploads: bool = True,
) -> dict:
    metadata = metadata or {}
    archive_id = existing_record["id"] if existing_record else str(uuid4())
    archive_uuid = UUID(archive_id)
    existing_materials = {
        item.get("path", "").lower(): item for item in (existing_record or {}).get("materials", [])
    }
    existing_documents = (existing_record or {}).get("_documents", {})
    normalized_uploads = {normalize_archive_path(path): data for path, data in uploads}
    basename_counts = Counter(Path(path).name for path in normalized_uploads)
    basename_uploads = {Path(path).name: data for path, data in normalized_uploads.items() if basename_counts[Path(path).name] == 1}
    manifest_basename_counts = Counter(Path(normalize_archive_path(item.path)).name for item in manifest)
    uploads_preserve_paths = any("/" in path for path in normalized_uploads)
    materials: list[dict] = []
    documents: dict[str, dict] = {}
    warnings: list[str] = []
    for source in manifest:
        path = normalize_archive_path(source.path)
        if not path:
            continue
        material_id = str(uuid5(archive_uuid, path.lower()))
        name = Path(path).name
        extension = Path(name).suffix.lower()
        item = {
            "id": material_id,
            "path": path,
            "name": name,
            "extension": extension,
            "size": source.size,
            "category": classify_material(path),
            "chapter": extract_chapter(path),
            "lesson": extract_lesson(path),
            "version": extract_version(path),
            "sha256": None,
            "duplicate_group": None,
            "parse_status": "metadata_only" if extension in PARSEABLE_EXTENSIONS else "unsupported",
            "parse_message": "未上传正文，仅完成文件名与路径分析" if extension in PARSEABLE_EXTENSIONS else "该格式仅纳入资料清单",
            "document_id": None,
            "preview_available": False,
            "character_count": 0,
            "excerpt": "",
            "last_modified": source.last_modified,
        }
        previous = existing_materials.get(path.lower())
        data = normalized_uploads.get(path)
        if data is None and not uploads_preserve_paths and basename_counts.get(name) == 1 and manifest_basename_counts.get(name) == 1:
            data = basename_uploads.get(name)
        if data is not None:
            item["sha256"] = hashlib.sha256(data).hexdigest()
            if previous and previous.get("sha256") == item["sha256"] and previous.get("document_id"):
                for key in ("parse_status", "parse_message", "document_id", "preview_available", "character_count", "excerpt"):
                    item[key] = previous.get(key, item[key])
                if previous["id"] in existing_documents:
                    documents[material_id] = existing_documents[previous["id"]]
            elif not extract_uploads:
                document_id = str(uuid4())
                persist_original(document_store_path, document_id, name, data)
                item.update({
                    "parse_status": "metadata_only" if extension in PARSEABLE_EXTENSIONS else "unsupported",
                    "parse_message": "原文件已保存；进入备课时再按需提取正文" if extension in PARSEABLE_EXTENSIONS else "原文件已保存；该格式暂不支持正文提取",
                    "document_id": document_id,
                    "preview_available": extension in PREVIEW_EXTENSIONS,
                })
            elif extension in PARSEABLE_EXTENSIONS:
                try:
                    parsed = _parse_uploaded(name, data, document_store_path)
                    documents[material_id] = parsed
                    item.update({
                        "parse_status": "parsed",
                        "parse_message": "正文、结构和候选知识点已提取",
                        "document_id": parsed["document_id"],
                        "preview_available": extension in PREVIEW_EXTENSIONS,
                        "character_count": parsed["character_count"],
                        "excerpt": re.sub(r"\s+", " ", parsed["raw_text"][:500]).strip(),
                    })
                except Exception as exc:
                    item["parse_status"] = "parse_failed"
                    item["parse_message"] = str(exc)[:240]
                    warnings.append(f"{path}：{item['parse_message']}")
        elif (
            previous
            and previous.get("size") == source.size
            and previous.get("last_modified") in {None, source.last_modified}
            and previous.get("document_id")
        ):
            for key in ("sha256", "parse_status", "parse_message", "document_id", "preview_available", "character_count", "excerpt"):
                item[key] = previous.get(key, item[key])
            if previous["id"] in existing_documents:
                documents[material_id] = existing_documents[previous["id"]]
        materials.append(item)
    duplicate_groups = _assign_duplicates(materials)
    chapters = _chapters(materials)
    schedule = _schedule_entries(materials, documents)
    habits = infer_preparation_habits(materials, documents)
    timestamp = utc_now()
    categories = dict(Counter(item["category"] for item in materials))
    display_name = archive_name.strip() or (manifest[0].path.split("/")[0] if manifest else "学期资料库")
    record = {
        "id": archive_id,
        "name": display_name[:100],
        "course_title": metadata.get("course_title") or (existing_record or {}).get("course_title") or _course_title(display_name, materials, documents),
        "academic_term": metadata.get("academic_term", (existing_record or {}).get("academic_term", "")),
        "course_code": metadata.get("course_code", (existing_record or {}).get("course_code", "")),
        "local_root": metadata.get("local_root", (existing_record or {}).get("local_root")),
        "last_scanned_at": metadata.get("last_scanned_at", (existing_record or {}).get("last_scanned_at")),
        "source_folders": (existing_record or {}).get("source_folders", []),
        "created_at": (existing_record or {}).get("created_at", timestamp),
        "updated_at": timestamp,
        "total_files": len(materials),
        "parsed_files": len(documents),
        "duplicate_groups": duplicate_groups,
        "chapter_count": len(chapters),
        "categories": categories,
        "materials": materials,
        "chapters": chapters,
        "schedule": schedule,
        "habits": habits,
        "preparation_profile_prompt": _profile_prompt(habits),
        "warnings": warnings[:30],
        "_documents": documents,
    }
    return record


def extract_course_archive_materials(
    record: dict,
    material_ids: list[str],
    document_store_path: Path,
    *,
    engine: str = "mineru",
    force: bool = False,
    progress_cb: Callable[[int, str], None] | None = None,
) -> dict:
    requested = set(material_ids)
    materials = {item["id"]: item for item in record.get("materials", [])}
    missing = requested.difference(materials)
    if missing:
        raise KeyError("部分待提取资料已不存在，请刷新后重新选择")

    documents = record.setdefault("_documents", {})
    sources = {item["id"]: item for item in record.get("source_folders", [])}
    errors: list[str] = []
    for material_id in material_ids:
        material = materials[material_id]
        if material.get("parse_status") == "parsed" and material_id in documents and not force:
            continue
        if material.get("extension") not in PARSEABLE_EXTENSIONS:
            errors.append(f"{material['name']}：该格式暂不支持正文提取")
            continue
        try:
            if not material.get("document_id"):
                source = sources.get(material.get("source_folder_id"))
                if not source or source.get("kind") != "local" or not source.get("root_path"):
                    errors.append(f"{material['name']}：尚未导入原件，请在数据中台刷新该浏览器来源后重试")
                    continue
                root = Path(source["root_path"]).expanduser().resolve()
                relative = material.get("source_relative_path") or material["name"]
                local_file = (root / relative).resolve()
                if root not in local_file.parents or not local_file.is_file():
                    errors.append(f"{material['name']}：本地源文件已移动或不存在，请刷新来源目录")
                    continue
                local_data = local_file.read_bytes()
                document_id = str(uuid4())
                persist_original(document_store_path, document_id, material["name"], local_data)
                material.update({
                    "document_id": document_id,
                    "sha256": hashlib.sha256(local_data).hexdigest(),
                    "preview_available": material.get("extension") in PREVIEW_EXTENSIONS,
                    "parse_message": "已按本次备课需要读取本地原件",
                })
            source = original_path(document_store_path, material["document_id"])
            parsed = _parse_uploaded(
                material["name"], source.read_bytes(), document_store_path, material["document_id"],
                full_extraction=True, engine=engine, progress_cb=progress_cb,
            )
            documents[material_id] = parsed
            material.update({
                "parse_status": "parsed",
                "parse_message": "已按本次备课需要提取正文、结构和候选知识点",
                "preview_available": material.get("extension") in PREVIEW_EXTENSIONS,
                "character_count": parsed["character_count"],
                "excerpt": re.sub(r"\s+", " ", parsed["raw_text"][:500]).strip(),
            })
        except Exception as exc:
            material["parse_status"] = "parse_failed"
            material["parse_message"] = str(exc)[:240]
            errors.append(f"{material['name']}：{material['parse_message']}")

    if errors:
        record["warnings"] = list(dict.fromkeys([*record.get("warnings", []), *errors]))[-30:]
    record["parsed_files"] = len(documents)
    record["schedule"] = _schedule_entries(record.get("materials", []), documents)
    record["habits"] = infer_preparation_habits(record.get("materials", []), documents)
    record["preparation_profile_prompt"] = _profile_prompt(record["habits"])
    record["updated_at"] = utc_now()
    return record


def append_course_archive_files(
    record: dict,
    uploads: list[tuple[str, bytes]],
    document_store_path: Path,
    chapter: str | None,
) -> tuple[dict, list[dict]]:
    existing_paths = {item.get("path", "").casefold() for item in record.get("materials", [])}
    normalized_uploads: list[tuple[str, bytes]] = []
    for raw_path, data in uploads:
        path = normalize_archive_path(raw_path) or "未命名资料"
        candidate = path
        index = 2
        while candidate.casefold() in existing_paths:
            source = PurePosixPath(path)
            candidate = str(source.with_name(f"{source.stem} ({index}){source.suffix}"))
            index += 1
        existing_paths.add(candidate.casefold())
        normalized_uploads.append((candidate, data))

    manifest = [
        ArchiveManifestItem(path=item["path"], size=item.get("size", 0), last_modified=item.get("last_modified"))
        for item in record.get("materials", [])
    ]
    manifest.extend(ArchiveManifestItem(path=path, size=len(data)) for path, data in normalized_uploads)
    updated = analyze_course_archive(
        record.get("name", "课程资料库"),
        manifest,
        normalized_uploads,
        document_store_path,
        record,
        None,
        False,
    )
    new_paths = {path.casefold() for path, _ in normalized_uploads}
    added = [item for item in updated.get("materials", []) if item.get("path", "").casefold() in new_paths]
    if chapter:
        for material in added:
            material["chapter"] = chapter
        updated["chapters"] = _chapters(updated.get("materials", []))
        updated["chapter_count"] = len(updated["chapters"])
    updated["updated_at"] = utc_now()
    return updated, added


def remove_course_archive_materials(record: dict, material_ids: set[str]) -> tuple[dict, list[str]]:
    removed = [item for item in record.get("materials", []) if item["id"] in material_ids]
    if not removed:
        raise KeyError("未找到需要删除的原始文件")
    record["materials"] = [item for item in record.get("materials", []) if item["id"] not in material_ids]
    documents = record.setdefault("_documents", {})
    for material_id in material_ids:
        documents.pop(material_id, None)
    record["total_files"] = len(record["materials"])
    record["parsed_files"] = len(documents)
    record["duplicate_groups"] = _assign_duplicates(record["materials"])
    record["chapters"] = _chapters(record["materials"])
    record["chapter_count"] = len(record["chapters"])
    record["categories"] = dict(Counter(item["category"] for item in record["materials"]))
    record["schedule"] = _schedule_entries(record["materials"], documents)
    record["habits"] = infer_preparation_habits(record["materials"], documents)
    record["preparation_profile_prompt"] = _profile_prompt(record["habits"])
    record["updated_at"] = utc_now()
    return record, [item["document_id"] for item in removed if item.get("document_id")]


def archive_summary(record: dict) -> dict:
    return _summary(record)


def prepare_archive_pack(record: dict, request: PrepareArchiveRequest) -> dict:
    materials_by_id = {item["id"]: item for item in record.get("materials", [])}
    documents = record.get("_documents", {})
    selected_ids = [item_id for item_id in request.material_ids if item_id in materials_by_id]
    if not selected_ids and request.chapter:
        selected_ids = [
            item["id"] for item in record.get("materials", [])
            if item.get("chapter") == request.chapter and item["id"] in documents
        ]
    if not selected_ids:
        selected_ids = list(documents)[:12]
    parsed_ids = [item_id for item_id in selected_ids if item_id in documents]
    if not parsed_ids:
        raise ValueError("所选资料尚未提取正文，请重新导入并确保至少上传一份 PDF、DOCX、PPTX 或文本材料")
    primary_id = request.primary_material_id if request.primary_material_id in parsed_ids else None
    if primary_id is None:
        priority = {"textbook": 0, "courseware": 1, "lesson_plan": 2, "syllabus": 3, "experiment": 4}
        primary_id = min(parsed_ids, key=lambda item_id: priority.get(materials_by_id[item_id]["category"], 9))
    primary = materials_by_id[primary_id]
    supporting = [item_id for item_id in parsed_ids if item_id != primary_id]
    resources = []
    for item_id in [primary_id, *supporting]:
        item = materials_by_id[item_id]
        resources.append({
            "id": item_id,
            "name": item["name"],
            "category": item["category"],
            "chapter": item.get("chapter"),
            "role": "primary" if item_id == primary_id else "supporting",
            "document_id": item.get("document_id"),
            "preview_available": item.get("preview_available", False),
        })
    context_sections = [record.get("preparation_profile_prompt", "")]
    if request.chapter:
        context_sections.append(f"# 当前教学范围\n章节：{request.chapter}")
    if request.session_label:
        context_sections.append(f"# 当前讲次\n{request.session_label}")
    excerpts: list[str] = []
    # Workflow context accepts 10,000 characters. Keep room for teacher notes and
    # optional visual-page evidence added by the frontend at launch time.
    remaining = 5000
    for item_id in supporting:
        item = materials_by_id[item_id]
        text = documents[item_id].get("raw_text", "")
        excerpt = text[: min(3500, remaining)].strip()
        if excerpt:
            excerpts.append(f"## {item['name']}（{item['category']}）\n{excerpt}")
            remaining -= len(excerpt)
        if remaining <= 0:
            break
    if excerpts:
        context_sections.append("# 配套资料证据\n" + "\n\n".join(excerpts))
    label = request.chapter or request.session_label
    title = record["course_title"] + (f" · {label}" if label else " · 当前讲次")
    context = "\n\n".join(section for section in context_sections if section)
    if len(context) > 7800:
        context = context[:7750].rstrip() + "\n\n[配套资料摘录已按工作流上下文上限截断]"
    return {
        "archive_id": record["id"],
        "title": title,
        "chapter": request.chapter,
        "session_label": request.session_label,
        "primary_material_id": primary_id,
        "parsed_document": documents[primary_id],
        "resources": resources,
        "context": context,
        "preparation_profile_prompt": record.get("preparation_profile_prompt", ""),
    }
