import re
from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4, uuid5, UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from backend.course_designs.models import (
    CourseDesignAssemblyApply, CourseDesignContent, CourseDesignCreate,
)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def public_record(record: dict) -> dict:
    return {key: value for key, value in record.items() if not key.startswith("_")}


def summary(record: dict) -> dict:
    exports = record.get("exports", [])
    return {
        "id": record["id"],
        "title": record["title"],
        "archive_id": record["archive_id"],
        "chapter": record.get("chapter"),
        "run_id": record.get("run_id"),
        "status": record.get("status", "draft"),
        "version": record.get("version", 1),
        "source_count": len(record.get("source_references", [])),
        "export_count": len(exports),
        "latest_export_at": exports[-1]["created_at"] if exports else None,
        "updated_at": record["updated_at"],
    }


def _ref_id(design_id: str, value: str) -> str:
    return str(uuid5(UUID(design_id), value))


def _reference(design_id: str, archive: dict, material: dict, layer: str) -> dict:
    document_id = material.get("document_id")
    suffix = f"/api/documents/{document_id}" if document_id else None
    locator = material.get("path", "") if layer == "original" else f"material:{material['id']}:extracted"
    return {
        "id": _ref_id(design_id, f"{layer}:{material['id']}"),
        "layer": layer,
        "archive_id": archive["id"],
        "material_id": material["id"],
        "document_id": document_id,
        "source_name": material["name"],
        "source_path": material.get("path", ""),
        "locator": locator,
        "sha256": material.get("sha256"),
        "extraction_status": material.get("parse_status", ""),
        "character_count": material.get("character_count", 0),
        "excerpt": material.get("excerpt", "") if layer == "extracted" else "",
        "original_url": f"{suffix}/original" if suffix else None,
        "preview_url": f"{suffix}/preview" if suffix and material.get("preview_available") else None,
    }


def create_design(archive: dict, payload: CourseDesignCreate, outline: dict | None = None) -> dict:
    materials = {item["id"]: item for item in archive.get("materials", [])}
    missing_ids = [item_id for item_id in payload.material_ids if item_id not in materials]
    if missing_ids:
        raise ValueError(f"所选资料不属于当前课程资料库：{', '.join(missing_ids)}")
    selected_ids = [item_id for item_id in payload.material_ids if item_id in materials]
    if not selected_ids:
        selected_ids = [
            item["id"] for item in archive.get("materials", [])
            if item.get("parse_status") == "parsed" and (not payload.chapter or item.get("chapter") == payload.chapter)
        ][:20]
    if not selected_ids:
        raise ValueError("当前范围没有可用于课程设计的已提取资料")
    primary_id = payload.primary_material_id if payload.primary_material_id in selected_ids else selected_ids[0]
    design_id = str(uuid4())
    references = []
    for item_id in selected_ids:
        references.append(_reference(design_id, archive, materials[item_id], "original"))
        if materials[item_id].get("parse_status") == "parsed":
            references.append(_reference(design_id, archive, materials[item_id], "extracted"))
    schedule = next((item for item in archive.get("schedule", []) if item["id"] == payload.schedule_id), None)
    selected_schedule_ids = {
        str(item_id).rsplit(":", 1)[-1]
        for item_id in (outline or {}).get("selected_session_ids", [])
    }
    selected_schedules = [
        deepcopy(item) for item in archive.get("schedule", [])
        if item.get("id") in selected_schedule_ids
    ]
    if schedule and schedule.get("id") not in {item.get("id") for item in selected_schedules}:
        selected_schedules.insert(0, deepcopy(schedule))
    if schedule:
        references.append({
            "id": _ref_id(design_id, f"schedule:{schedule['id']}"),
            "layer": "structured",
            "archive_id": archive["id"],
            "material_id": schedule.get("source_material_id"),
            "document_id": materials.get(schedule.get("source_material_id", ""), {}).get("document_id"),
            "source_name": schedule["label"],
            "source_path": materials.get(schedule.get("source_material_id", ""), {}).get("path", ""),
            "locator": f"schedule:{schedule['id']}",
            "sha256": materials.get(schedule.get("source_material_id", ""), {}).get("sha256"),
            "extraction_status": "structured",
            "character_count": len(schedule.get("content", "")),
            "excerpt": schedule.get("content", ""),
            "original_url": None,
            "preview_url": None,
        })
    if outline:
        outline_locator = (
            f"material-unit:{payload.material_unit_id}:knowledge-outline:"
            f"{outline['id']}:v{outline['version']}"
        )
        references.append({
            "id": _ref_id(design_id, outline_locator),
            "layer": "structured",
            "archive_id": archive["id"],
            "material_id": None,
            "document_id": None,
            "source_name": outline["title"],
            "source_path": "",
            "locator": outline_locator,
            "sha256": None,
            "extraction_status": "structured",
            "character_count": sum(
                len(item.get("title", "")) for item in outline.get("knowledge_nodes", [])
            ),
            "excerpt": "、".join(
                item.get("title", "") for item in outline.get("knowledge_nodes", [])[:8]
            ),
            "original_url": None,
            "preview_url": None,
        })
    primary_document = archive.get("_documents", {}).get(primary_id, {})
    points = primary_document.get("knowledge_points", [])
    point_titles = [item.get("title", "").strip() for item in points if item.get("title", "").strip()]
    key_points = [item.get("title", "").strip() for item in points if item.get("is_key_point") and item.get("title", "").strip()] or point_titles[:6]
    difficult = [item.get("title", "").strip() for item in points if item.get("difficulty_level") in {"困难", "较难", "高"} and item.get("title", "").strip()]
    if outline:
        outline_nodes = [
            item for item in outline.get("knowledge_nodes", [])
            if item.get("title", "").strip()
        ]
        outline_titles = [item["title"].strip() for item in outline_nodes]
        outline_key_points = [
            item["title"].strip() for item in outline_nodes if item.get("is_key_point")
        ]
        key_points = outline_key_points
        difficult = [
            item["title"].strip() for item in outline_nodes if item.get("is_difficult_point")
        ]
    requirements = list((outline or {}).get("requirements") or [])
    requirement_values: dict[str, list[str]] = {}
    for requirement in requirements:
        category = str(requirement.get("category") or "knowledge")
        value = str(requirement.get("content") or requirement.get("title") or "").strip()
        if value and value not in requirement_values.setdefault(category, []):
            requirement_values[category].append(value)
    chapter = payload.chapter or materials[primary_id].get("chapter") or ""
    session_label = (
        outline.get("session", "")
        if outline
        else (schedule.get("content", "") if schedule else "")
    )
    topic = (
        outline.get("title", "") if outline else session_label
    ) or "、".join(key_points[:3]) or chapter or archive["course_title"]
    difficult_points = difficult if outline else (difficult or key_points[1:3])
    objectives = requirement_values.get("objective", [])[:20] or [
        f"准确说明{key_points[0] if key_points else topic}的核心概念与适用边界",
        f"结合课程材料分析{key_points[1] if len(key_points) > 1 else topic}的原理和应用",
        "依据任务条件完成方案比较、论证与规范表达",
    ]
    if requirement_values.get("key_point"):
        key_points = list(dict.fromkeys([*requirement_values["key_point"], *key_points]))
    if requirement_values.get("difficult_point"):
        difficult_points = list(dict.fromkeys([*requirement_values["difficult_point"], *difficult_points]))
    assessment_requirements = requirement_values.get("assessment", [])
    practice_requirements = requirement_values.get("practice", [])
    content = CourseDesignContent(
        course_name=archive["course_title"],
        topic=topic[:240],
        chapter=chapter,
        session_label=session_label[:300],
        objectives=objectives,
        knowledge_points=(outline_titles if outline else point_titles)[:240],
        key_points=key_points[:12],
        difficult_points=difficult_points[:8],
        methods=["讲授与案例分析", "问题驱动", "对比归纳"],
        tools=["原始教材与课件", "工程案例或仿真材料"],
        assessment="\n".join(assessment_requirements) or "通过概念辨析、案例论证和课堂产出检查目标达成情况。",
        teaching_process=(
            "一、问题与材料定位\n二、核心概念和原理分析\n三、案例比较与方案论证\n"
            + ("实践要求：" + "；".join(practice_requirements) + "\n" if practice_requirements else "")
            + "四、归纳总结与形成性评价"
        ),
    ).model_dump()
    timestamp = utc_now()
    return {
        "id": design_id,
        "title": outline.get("title", "") if outline else f"{archive['course_title']} · {chapter or '课程设计'}",
        "archive_id": archive["id"],
        "chapter": chapter or None,
        "schedule_id": payload.schedule_id,
        "primary_material_id": primary_id,
        "material_ids": selected_ids,
        "material_unit_id": payload.material_unit_id,
        "knowledge_outline_id": outline.get("id") if outline else None,
        "knowledge_outline_version": outline.get("version") if outline else None,
        "run_id": None,
        "status": "draft",
        "version": 1,
        "template_document_id": None,
        "template_material_id": None,
        "source_snapshot": {
            "schedule": selected_schedules or ([{
                "id": item_id, "label": "选定讲次", "content": session_label,
            } for item_id in (outline or {}).get("selected_session_ids", [])] if session_label else []),
            "syllabus_requirements": deepcopy(requirements),
            "knowledge_nodes": deepcopy((outline or {}).get("knowledge_nodes", [])),
        },
        "content_insertions": [],
        "exports": [],
        "source_references": references,
        "content": content,
        "created_at": timestamp,
        "updated_at": timestamp,
        "_versions": [{
            "version": 1,
            "status": "draft",
            "content": deepcopy(content),
            "template_document_id": None,
            "template_material_id": None,
            "content_insertions": [],
            "created_at": timestamp,
        }],
    }


def rebind_design(record: dict, outline: dict | None, archive: dict) -> dict:
    """把课程设计重新绑定到资料单元大纲的最新版本(显式升级)。

    教案是交付物，不能大纲一变就悄悄改。只有教师点「一键升级」才调到本逻辑：
    用最新版大纲重烘焙 snapshot，并更新 content 的重难点/知识点/来源引用。
    """
    if outline is None:
        raise ValueError("资料单元尚未生成知识大纲，无法重新绑定")
    record["knowledge_outline_id"] = outline["id"]
    record["knowledge_outline_version"] = deepcopy(outline["version"])
    outline_nodes = [
        item for item in outline.get("knowledge_nodes") or [] if item.get("title", "").strip()
    ]
    outline_titles = [item["title"].strip() for item in outline_nodes]
    content = CourseDesignContent.model_validate(record["content"])
    content.topic = outline.get("title") or content.topic
    content.session_label = outline.get("session") or content.session_label
    content.knowledge_points = outline_titles[:240]
    content.key_points = [item["title"].strip() for item in outline_nodes if item.get("is_key_point")][:12]
    content.difficult_points = [item["title"].strip() for item in outline_nodes if item.get("is_difficult_point")][:8]
    # 保留教师已在教案里手动改过的字段不做覆盖：只更新大纲来源应负责的知识点/重难点
    requirements = outline.get("requirements") or []
    requirement_values: dict[str, list[str]] = {}
    for requirement in requirements:
        category = str(requirement.get("category") or "knowledge")
        value = str(requirement.get("content") or requirement.get("title") or "").strip()
        if value and value not in requirement_values.setdefault(category, []):
            requirement_values[category].append(value)
    if requirement_values.get("key_point"):
        content.key_points = list(dict.fromkeys([*requirement_values["key_point"], *content.key_points]))[:12]
    if requirement_values.get("difficult_point"):
        content.difficult_points = list(dict.fromkeys([*requirement_values["difficult_point"], *content.difficult_points]))[:8]
    record["source_snapshot"] = {
        "schedule": _schedule_snapshot(archive, outline, record.get("schedule_id")),
        "syllabus_requirements": deepcopy(requirements),
        "knowledge_nodes": deepcopy(outline.get("knowledge_nodes") or []),
    }
    # 更新 source_references 里大纲那一条的 locator：移除所有旧的 material-unit 大纲引用，替换为当前版本
    outline_locator = (
        f"material-unit:{record.get('material_unit_id')}:knowledge-outline:"
        f"{outline['id']}:v{outline['version']}"
    )
    ref_id = _ref_id(record["id"], outline_locator)
    record["source_references"] = [
        item for item in record.get("source_references", [])
        if not (
            str(item.get("locator") or "").startswith("material-unit:")
            and f":{outline['id']}:" in str(item.get("locator"))
        )
    ]
    record["source_references"].append({
        "id": ref_id,
        "layer": "structured",
        "archive_id": record["archive_id"],
        "material_id": None,
        "document_id": None,
        "source_name": outline["title"],
        "source_path": "",
        "locator": outline_locator,
        "sha256": None,
        "extraction_status": "structured",
        "character_count": sum(len(item.get("title", "")) for item in outline.get("knowledge_nodes", [])),
        "excerpt": "、".join(item.get("title", "") for item in outline.get("knowledge_nodes", [])[:8]),
        "original_url": None,
        "preview_url": None,
    })
    # 升级 = 一次内容修订：version+1 并保留 _versions 历史，与教师手改一致，可回溯
    record["content"] = content.model_dump()
    next_version = record.get("version", 1) + 1
    record["version"] = next_version
    timestamp = utc_now()
    record.setdefault("_versions", []).append({
        "version": next_version,
        "status": record.get("status", "draft"),
        "content": content.model_dump(),
        "template_document_id": record.get("template_document_id"),
        "template_material_id": record.get("template_material_id"),
        "content_insertions": deepcopy(record.get("content_insertions") or []),
        "created_at": timestamp,
    })
    record["_versions"] = record["_versions"][-30:]
    record["updated_at"] = timestamp
    return record


def _schedule_snapshot(archive: dict, outline: dict | None, schedule_id: str | None) -> list[dict]:
    """重建 source_snapshot.schedule —— 与 create_design 的 selected_schedules 逻辑一致。"""
    if not outline and not schedule_id:
        return []
    selected_schedule_ids = {
        str(item_id).rsplit(":", 1)[-1]
        for item_id in (outline or {}).get("selected_session_ids", [])
    }
    if schedule_id:
        selected_schedule_ids.add(str(schedule_id))
    schedules = [
        deepcopy(item) for item in archive.get("schedule") or []
        if item.get("id") in selected_schedule_ids
    ]
    return schedules


def update_design(
    record: dict,
    content: CourseDesignContent,
    status: str,
    template_document_id: str | None,
    template_material_id: str | None = None,
) -> dict:
    next_version = record.get("version", 1) + 1
    timestamp = utc_now()
    record.update({
        "content": content.model_dump(),
        "status": status,
        "version": next_version,
        "template_document_id": template_document_id,
        "template_material_id": template_material_id,
        "content_insertions": deepcopy(record.get("content_insertions", [])),
        "updated_at": timestamp,
    })
    record.setdefault("_versions", []).append({
        "version": next_version,
        "status": status,
        "content": content.model_dump(),
        "template_document_id": template_document_id,
        "template_material_id": template_material_id,
        "created_at": timestamp,
    })
    record["_versions"] = record["_versions"][-30:]
    return record


def validate_run_context(record: dict, run: dict) -> None:
    """Reject run/design links that cannot be proven to share the same source context."""
    teaching_data = run.get("teaching_data", {}) or {}
    run_archive_id = teaching_data.get("archive_id")
    run_design_id = teaching_data.get("design_id")
    run_document_id = teaching_data.get("document_id")
    source_document_ids = {
        item.get("document_id")
        for item in record.get("source_references", [])
        if item.get("document_id")
    }

    if run_design_id and run_design_id != record.get("id"):
        raise ValueError("生成会话属于其他课程设计，不能关联到当前教案")
    if run_archive_id and run_archive_id != record.get("archive_id"):
        raise ValueError("生成会话属于其他课程资料库，不能关联到当前教案")
    if run_document_id and run_document_id not in source_document_ids:
        raise ValueError("生成会话使用的主文档不在当前课程设计资料范围内")
    if not run_design_id and not run_archive_id and not run_document_id:
        raise ValueError("生成会话缺少课程、设计和文档来源，无法确认数据归属")


def sync_run(record: dict, run: dict, teacher_draft: str | None = None) -> dict:
    validate_run_context(record, run)
    content = CourseDesignContent.model_validate(record["content"])
    teaching = run.get("teaching_data", {})
    analysis = teaching.get("content_analysis", {}) or {}
    framework = teaching.get("teaching_framework", {}) or {}
    if framework.get("learning_objectives"):
        content.objectives = framework["learning_objectives"][:20]
    if analysis.get("key_points"):
        content.key_points = analysis["key_points"][:30]
    if analysis.get("difficult_points"):
        content.difficult_points = analysis["difficult_points"][:20]
    if framework.get("strategies"):
        content.methods = framework["strategies"][:20]
    stages = framework.get("stages", [])
    if stages:
        content.teaching_process = "\n\n".join(
            f"{index}. {stage.get('name', '教学环节')}\n{stage.get('purpose', '')}\n{stage.get('activity', '')}".strip()
            for index, stage in enumerate(stages, 1)
        )
    if framework.get("assessment"):
        content.assessment = "\n".join(framework["assessment"])
    ideological = framework.get("ideological_elements") or []
    if isinstance(ideological, list):
        values = []
        for item in ideological:
            if isinstance(item, dict):
                value = "；".join(
                    str(item.get(key) or "").strip()
                    for key in ("dimension", "content", "integration_method")
                    if str(item.get(key) or "").strip()
                )
            else:
                value = str(item).strip()
            if value and value not in values:
                values.append(value)
        if values:
            content.ideological_elements = values[:20]
    record["run_id"] = run["id"]
    generated_id = _ref_id(record["id"], f"run:{run['id']}")
    record["source_references"] = [item for item in record.get("source_references", []) if item["id"] != generated_id]
    record["source_references"].append({
        "id": generated_id,
        "layer": "generated",
        "archive_id": record["archive_id"],
        "source_name": run.get("objective", "多智能体课程设计"),
        "source_path": "",
        "locator": f"workflow-run:{run['id']}",
        "extraction_status": run.get("status", ""),
        "character_count": len(teacher_draft or run.get("final_output") or ""),
        "excerpt": (teacher_draft or run.get("final_output") or "")[:500],
    })
    return update_design(
        record,
        content,
        record.get("status", "draft"),
        record.get("template_document_id"),
        record.get("template_material_id"),
    )


_LIST_TARGETS = {
    "objectives", "knowledge_points", "key_points", "difficult_points", "methods", "tools",
    "ideological_elements",
}


def _source_item(
    source_id: str,
    kind: str,
    title: str,
    content: str,
    target: str,
    *,
    source_name: str = "",
    locator: str = "",
    iteration: int | None = None,
    category: str = "",
) -> dict:
    return {
        "id": source_id, "kind": kind, "title": title, "content": content.strip(),
        "source_name": source_name or title, "locator": locator, "default_target": target,
        "iteration": iteration, "category": category,
    }


def assembly_sources(
    record: dict,
    run: dict | None = None,
    teacher_draft: str | None = None,
) -> list[dict]:
    """Build traceable, user-selectable content blocks for the final lesson plan."""
    items: list[dict] = []
    snapshot = record.get("source_snapshot") or {}
    for index, schedule in enumerate(snapshot.get("schedule") or []):
        content = str(schedule.get("content") or schedule.get("label") or "").strip()
        if content:
            items.append(_source_item(
                f"schedule:{schedule.get('id') or index}", "schedule",
                str(schedule.get("label") or "进度表讲次"), content, "session_label",
                source_name="教学进度表", locator=f"schedule:{schedule.get('id') or index}",
            ))
    category_targets = {
        "objective": "objectives", "knowledge": "knowledge_points", "key_point": "key_points",
        "difficult_point": "difficult_points", "practice": "teaching_process", "assessment": "assessment",
    }
    for index, requirement in enumerate(snapshot.get("syllabus_requirements") or []):
        content = str(requirement.get("content") or requirement.get("title") or "").strip()
        if not content:
            continue
        category = str(requirement.get("category") or "knowledge")
        evidence = requirement.get("evidence") or {}
        items.append(_source_item(
            f"syllabus:{requirement.get('id') or index}", "syllabus",
            str(requirement.get("category_label") or requirement.get("title") or "教学大纲要求"),
            content, category_targets.get(category, "knowledge_points"),
            source_name=str(evidence.get("label") or "课程教学大纲"),
            locator=str(evidence.get("locator") or f"syllabus:{index}"), category=category,
        ))
    for index, node in enumerate(snapshot.get("knowledge_nodes") or []):
        title = str(node.get("title") or "").strip()
        description = str(node.get("description") or "").strip()
        content = f"{title}{'：' + description if description else ''}"
        if title:
            items.append(_source_item(
                f"outline:{node.get('id') or index}", "knowledge_outline", title, content,
                "difficult_points" if node.get("is_difficult_point") else "key_points" if node.get("is_key_point") else "knowledge_points",
                source_name="已确认知识大纲", locator=f"knowledge-outline:{node.get('id') or index}",
            ))
        # 教学补充作为独立可插入块——把资料单元生成的教学补充带进教案组装
        teacher_note = str(node.get("teacher_note") or "").strip()
        if teacher_note:
            items.append(_source_item(
                f"outline:{node.get('id') or index}:note", "knowledge_outline",
                f"{title} · 教学补充", teacher_note, "postscript",
                source_name="知识大纲教学补充", locator=f"knowledge-outline:{node.get('id') or index}:note",
            ))
    if run:
        for message in (run.get("teaching_data") or {}).get("messages") or []:
            if message.get("agent_type") != "teacher" or message.get("phase") not in {"design", "teach_knowledge", "teacher_answer"}:
                continue
            content = str(message.get("content") or "").strip()
            if not content:
                continue
            phase = str(message.get("phase"))
            iteration = int(message.get("iteration") or 0)
            phase_label = {"design": "教学设计", "teach_knowledge": "教师讲授", "teacher_answer": "教师答疑"}[phase]
            items.append(_source_item(
                f"teacher-message:{message.get('id')}", "teacher_message",
                f"{phase_label} · {'准备阶段' if iteration == 0 else f'第 {iteration} 轮'}", content,
                "teaching_process", source_name=str(message.get("agent_name") or "课程教师"),
                locator=f"workflow-run:{run.get('id')}:{phase}:{message.get('id')}", iteration=iteration,
            ))
        framework = (run.get("teaching_data") or {}).get("teaching_framework") or {}
        ideology = framework.get("ideological_elements") or framework.get("ideological_mapping") or []
        if isinstance(ideology, dict):
            ideology = ["：".join(filter(None, (str(key), str(value)))) for key, value in ideology.items()]
        if isinstance(ideology, str):
            ideology = [ideology]
        for index, value in enumerate(ideology):
            if isinstance(value, dict):
                value = "；".join(str(item) for item in value.values() if item)
            if str(value).strip():
                items.append(_source_item(
                    f"ideological:{index}", "ideological", f"课程思政建议 {index + 1}", str(value),
                    "ideological_elements", source_name="课程教师", locator=f"workflow-run:{run.get('id')}:ideological:{index}",
                ))
        if teacher_draft and teacher_draft.strip():
            items.append(_source_item(
                f"teacher-draft:{run.get('id')}", "teacher_draft", "教师审核稿（完整）",
                teacher_draft, "teaching_process", source_name="教师审核稿",
                locator=f"workflow-run:{run.get('id')}:teacher-draft",
            ))
    return items


def restore_source_snapshot(record: dict, archive: dict, unit: dict | None) -> dict:
    """Restore snapshots for records created before result assembly was introduced."""
    snapshot = record.get("source_snapshot") or {}
    if any(snapshot.get(key) for key in ("schedule", "syllabus_requirements", "knowledge_nodes")):
        return record
    outline = None
    if unit and record.get("knowledge_outline_id"):
        candidates = [
            item for item in unit.get("knowledge_outlines") or []
            if item.get("id") == record.get("knowledge_outline_id")
            and (
                record.get("knowledge_outline_version") is None
                or int(item.get("version") or 0) == int(record["knowledge_outline_version"])
            )
        ]
        outline = max(candidates, key=lambda item: int(item.get("version") or 0)) if candidates else None
    selected_schedule_ids = {
        str(item_id).rsplit(":", 1)[-1]
        for item_id in (outline or {}).get("selected_session_ids", [])
    }
    if record.get("schedule_id"):
        selected_schedule_ids.add(str(record["schedule_id"]))
    schedules = [
        deepcopy(item) for item in archive.get("schedule") or []
        if item.get("id") in selected_schedule_ids
    ]
    if not schedules and record.get("content", {}).get("session_label"):
        schedules = [{
            "id": record.get("schedule_id") or "legacy-session",
            "label": "已选讲次", "content": record["content"]["session_label"],
        }]
    knowledge_nodes = deepcopy((outline or {}).get("nodes") or (outline or {}).get("knowledge_nodes") or [])
    if not knowledge_nodes:
        content = record.get("content") or {}
        key_points = set(content.get("key_points") or [])
        difficult_points = set(content.get("difficult_points") or [])
        knowledge_nodes = [{
            "id": f"legacy-{index}", "title": title, "level": 1,
            "description": "", "is_key_point": title in key_points,
            "is_difficult_point": title in difficult_points,
        } for index, title in enumerate(content.get("knowledge_points") or []) if str(title).strip()]
    record["source_snapshot"] = {
        "schedule": schedules,
        "syllabus_requirements": deepcopy((outline or {}).get("requirements") or []),
        "knowledge_nodes": knowledge_nodes,
    }
    record.setdefault("content_insertions", [])
    return record


def apply_assembly(
    record: dict,
    payload: CourseDesignAssemblyApply,
    available_sources: list[dict],
) -> dict:
    source_map = {item["id"]: item for item in available_sources}
    missing = [source_id for source_id in payload.source_ids if source_id not in source_map]
    if missing:
        raise KeyError("部分插入来源已失效，请刷新内容清单后重试")
    selected = [source_map[source_id] for source_id in payload.source_ids]
    if payload.custom_content.strip():
        selected.append(_source_item(
            f"custom:{uuid4()}", "custom", payload.custom_title, payload.custom_content,
            payload.target_field, source_name="教师填写", locator="course-design:custom",
        ))
    content = CourseDesignContent.model_validate(record["content"])
    target = payload.target_field
    incoming_parts = [item["content"].strip() for item in selected if item["content"].strip()]
    if target in _LIST_TARGETS:
        incoming: list[str] = []
        for part in incoming_parts:
            lines = [re.sub(r"^\s*[-*\d.、]+\s*", "", line).strip() for line in part.splitlines()]
            incoming.extend(line for line in lines if line and line not in incoming)
        current = list(getattr(content, target))
        if payload.mode == "replace":
            value = incoming
        elif payload.mode == "prepend":
            value = list(dict.fromkeys([*incoming, *current]))
        else:
            value = list(dict.fromkeys([*current, *incoming]))
        setattr(content, target, value)
    else:
        incoming_text = "\n\n".join(incoming_parts)
        current_text = str(getattr(content, target) or "").strip()
        if payload.mode == "replace":
            value = incoming_text
        elif payload.mode == "prepend":
            value = "\n\n".join(filter(None, (incoming_text, current_text)))
        else:
            value = "\n\n".join(filter(None, (current_text, incoming_text)))
        setattr(content, target, value)
    now = utc_now()
    insertions = list(record.get("content_insertions") or [])
    insertions.extend({
        "id": str(uuid4()), "source_id": item["id"], "source_kind": item["kind"],
        "source_name": item["source_name"], "locator": item.get("locator", ""),
        "target_field": target, "mode": payload.mode, "content_preview": item["content"][:500],
        "created_at": now,
    } for item in selected)
    record["content_insertions"] = insertions[-200:]
    return update_design(
        record, content, record.get("status", "draft"),
        record.get("template_document_id"), record.get("template_material_id"),
    )


def reference_detail(record: dict, archive: dict, reference_id: str) -> dict:
    reference = next((item for item in record.get("source_references", []) if item["id"] == reference_id), None)
    if not reference:
        raise KeyError("未找到数据引用")
    content = reference.get("excerpt", "")
    sections = []
    if reference.get("layer") == "extracted" and reference.get("material_id"):
        document = archive.get("_documents", {}).get(reference["material_id"], {})
        content = document.get("raw_text", "")
        sections = document.get("sections", [])
    return {"reference": reference, "content": content, "sections": sections}


def _set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def _replace_cell(cell, value: str) -> None:
    row_element = cell._tc.getparent()
    row_properties = row_element.find(qn("w:trPr"))
    if row_properties is not None:
        for height in list(row_properties.findall(qn("w:trHeight"))):
            row_properties.remove(height)
    template_paragraph = cell.paragraphs[0]
    style = template_paragraph.style
    alignment = template_paragraph.alignment
    run_properties = deepcopy(template_paragraph.runs[0]._r.rPr) if template_paragraph.runs and template_paragraph.runs[0]._r.rPr is not None else None
    template_paragraph.clear()
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    lines = value.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = template_paragraph if index == 0 else cell.add_paragraph()
        paragraph.style = style
        paragraph.alignment = alignment
        run = paragraph.add_run(line)
        if run_properties is not None:
            run._r.insert(0, deepcopy(run_properties))


def _replace_paragraph(paragraph, value: str) -> None:
    run_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(value)
    if run_properties is not None:
        run._r.insert(0, deepcopy(run_properties))


def _document_containers(document) -> tuple[list, list]:
    paragraphs = list(document.paragraphs)
    tables = list(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            paragraphs.extend(container.paragraphs)
            tables.extend(container.tables)
    seen_cells: set[int] = set()
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                paragraphs.extend(cell.paragraphs)
    return paragraphs, tables


def _unique_table_cells(table):
    seen: set[int] = set()
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell_key = id(cell._tc)
            if cell_key in seen:
                continue
            seen.add(cell_key)
            yield row_index, column_index, row, cell


def _container_has_content(container) -> bool:
    return bool(container.tables or any(paragraph.text.strip() for paragraph in container.paragraphs))


def _template_values(content: CourseDesignContent) -> dict[str, str]:
    return {
        "course_name": content.course_name,
        "topic": content.topic,
        "chapter": content.chapter,
        "chapter_content": content.chapter,
        "session_label": content.session_label,
        "class_name": content.class_name,
        "location": content.location,
        "hours": content.hours,
        "duration": content.hours,
        "objectives": "\n".join(content.objectives),
        "knowledge_points": "\n".join(f"{i}. {item}" for i, item in enumerate(content.knowledge_points, 1)),
        "key_points": "\n".join(f"{i}. {item}" for i, item in enumerate(content.key_points, 1)),
        "difficult_points": "\n".join(f"{i}. {item}" for i, item in enumerate(content.difficult_points, 1)),
        "methods": "、".join(content.methods),
        "teaching_methods": "、".join(content.methods),
        "tools": "、".join(content.tools),
        "teaching_tools": "、".join(content.tools),
        "ideological_elements": "；".join(content.ideological_elements),
        "teaching_process": content.teaching_process,
        "r_process": content.teaching_process,
        "assessment": content.assessment,
        "postscript": content.postscript,
        "teaching_notes": content.postscript,
    }


_FIELD_ALIASES = {
    "chapter_content": "chapter", "duration": "hours", "teaching_methods": "methods",
    "teaching_tools": "tools", "r_process": "teaching_process", "teaching_notes": "postscript",
}
_FIELD_TITLES = {
    "session_label": "讲次与授课范围", "objectives": "教学目标", "knowledge_points": "知识点",
    "key_points": "教学重点", "difficult_points": "教学难点", "methods": "教学方法",
    "tools": "教学手段", "ideological_elements": "课程思政", "teaching_process": "教学过程",
    "assessment": "评价设计", "postscript": "教学后记",
}


_RIGHT_TEMPLATE_FIELDS = {
    "课程名称": "course_name", "授课主题": "topic", "章节内容": "chapter",
    "授课章节": "chapter", "授课班级": "class_name", "上课地点": "location",
    "授课地点": "location", "地点": "location", "课时": "hours", "学时": "hours",
    "课程学时": "hours", "讲次": "session_label", "授课进度": "session_label", "教学进度": "session_label", "教学目的": "objectives",
    "教学目的与要求": "objectives", "教学目的要求": "objectives", "教学目标": "objectives", "知识点": "knowledge_points",
    "课程目标": "objectives", "教学内容": "knowledge_points", "授课内容": "knowledge_points",
    "教学重点": "key_points", "教学难点": "difficult_points", "教学方法": "methods",
    "教学手段": "tools",
}
_BELOW_TEMPLATE_FIELDS = {
    "教学过程": "teaching_process", "教学内容与过程": "teaching_process",
    "教学设计": "teaching_process", "教学设计与实施": "teaching_process", "教学活动设计": "teaching_process",
    "课程思政元素": "ideological_elements", "课程思政": "ideological_elements",
    "思政元素": "ideological_elements", "思政融入": "ideological_elements", "思政融入点": "ideological_elements",
    "课程思政教学设计": "ideological_elements",
    "形成性评价": "assessment", "教学评价": "assessment", "评价设计": "assessment",
    "教学后记": "postscript", "教后记": "postscript", "课后反思": "postscript", "教学反思": "postscript",
}


def _normalized_template_label(value: str) -> str:
    return re.sub(r"[\s:：()（）]", "", value or "")


def _template_field(label: str, fields: dict[str, str]) -> str | None:
    normalized = _normalized_template_label(label)
    return next((field for keyword, field in fields.items() if normalized == _normalized_template_label(keyword)), None)


def _right_cell(row, cell) -> object | None:
    unique: list[object] = []
    seen: set[int] = set()
    for candidate in row.cells:
        key = id(candidate._tc)
        if key not in seen:
            seen.add(key)
            unique.append(candidate)
    for index, candidate in enumerate(unique):
        if candidate._tc is cell._tc:
            return unique[index + 1] if index + 1 < len(unique) else None
    return None


def _below_cell(table, row_index: int, column_index: int) -> object | None:
    if row_index + 1 >= len(table.rows):
        return None
    cells = table.rows[row_index + 1].cells
    return cells[min(column_index, len(cells) - 1)] if cells else None


def _field_target(table, row_index: int, column_index: int, row, cell, field: str) -> object | None:
    right = _right_cell(row, cell)
    below = _below_cell(table, row_index, column_index)
    if field in _RIGHT_TEMPLATE_FIELDS.values():
        return right or below
    return below or right


def _combined_target_value(entries: list[tuple[str, str]]) -> str:
    unique: list[tuple[str, str]] = []
    seen: set[str] = set()
    for field, value in entries:
        if not value or field in seen:
            continue
        seen.add(field)
        unique.append((field, value))
    if len(unique) == 1:
        return unique[0][1]
    return "\n\n".join(f"【{_FIELD_TITLES.get(field, field)}】\n{value}" for field, value in unique)


def inspect_docx_template(template_path: Path, content: CourseDesignContent) -> dict:
    document = Document(template_path)
    values = _template_values(content)
    paragraphs, tables = _document_containers(document)
    token = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")
    matched_fields: set[str] = set()
    replacement_count = 0
    for paragraph in paragraphs:
        matches = [item for item in token.findall(paragraph.text) if item in values]
        matches = [_FIELD_ALIASES.get(item, item) for item in matches]
        matched_fields.update(matches)
        replacement_count += len(matches)
    for table in tables:
        for row_index, column_index, row, cell in _unique_table_cells(table):
            field = _template_field(cell.text, _RIGHT_TEMPLATE_FIELDS) or _template_field(cell.text, _BELOW_TEMPLATE_FIELDS)
            if field and values[field] and _field_target(table, row_index, column_index, row, cell, field) is not None:
                matched_fields.add(field); replacement_count += 1
    content_fields = {
        field for field in (
            "course_name", "topic", "chapter", "session_label", "class_name", "location", "hours",
            "objectives", "knowledge_points", "key_points", "difficult_points", "methods", "tools",
            "ideological_elements", "teaching_process", "assessment", "postscript",
        ) if values.get(field)
    }
    required_body = {"objectives", "knowledge_points", "key_points", "difficult_points", "teaching_process", "assessment"}
    compatible = bool(matched_fields & required_body)
    return {
        "template_mode": "source-template",
        "compatible": compatible,
        "matched_fields": sorted(matched_fields),
        "unmatched_fields": sorted(content_fields - matched_fields),
        "replacement_count": replacement_count,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "header_count": sum(_container_has_content(section.header) for section in document.sections),
        "footer_count": sum(_container_has_content(section.footer) for section in document.sections),
        "message": (
            (
                f"模板可保持原格式填充，识别到 {len(matched_fields)} 类课程设计字段；"
                f"其余 {len(content_fields - matched_fields)} 类非空内容将追加到模板末尾"
            )
            if compatible else "模板未识别到教学目标、知识点、重难点、教学过程或评价等正文槽位"
        ),
    }


def _render_template(template_path: Path, content: CourseDesignContent) -> tuple[bytes, dict]:
    document = Document(template_path)
    report = inspect_docx_template(template_path, content)
    values = _template_values(content)
    token = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")
    paragraphs, tables = _document_containers(document)
    for paragraph in paragraphs:
        text = paragraph.text
        replaced = token.sub(lambda match: values.get(match.group(1), match.group(0)), text)
        if replaced != text:
            _replace_paragraph(paragraph, replaced)
    target_values: dict[int, tuple[object, list[tuple[str, str]]]] = {}
    for table in tables:
        for row_index, column_index, row, cell in _unique_table_cells(table):
            field = _template_field(cell.text, _RIGHT_TEMPLATE_FIELDS) or _template_field(cell.text, _BELOW_TEMPLATE_FIELDS)
            if not field or not values[field]:
                continue
            target = _field_target(table, row_index, column_index, row, cell, field)
            if target is None:
                continue
            key = id(target._tc)
            if key not in target_values:
                target_values[key] = (target, [])
            target_values[key][1].append((field, values[field]))
    for target, entries in target_values.values():
        _replace_cell(target, _combined_target_value(entries))
    appended_fields = [
        field for field in report["unmatched_fields"]
        if field in _FIELD_TITLES and values.get(field)
    ]
    if appended_fields:
        heading = document.add_paragraph()
        _set_run_font(heading.add_run("平台补充内容"), 12, True)
        for field in appended_fields:
            paragraph = document.add_paragraph()
            _set_run_font(paragraph.add_run(_FIELD_TITLES[field]), 11, True)
            for line in values[field].splitlines() or [values[field]]:
                paragraph = document.add_paragraph()
                _set_run_font(paragraph.add_run(line))
    output = BytesIO()
    document.save(output)
    return output.getvalue(), report


def _standard_document(content: CourseDesignContent, references: list[dict]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.add_run("课程教学设计"), 18, True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(subtitle.add_run(f"{content.course_name}  {content.topic}"), 12, True)
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    metadata = [("课程名称", content.course_name), ("章节", content.chapter), ("授课主题", content.topic), ("课时", content.hours), ("授课班级", content.class_name), ("地点", content.location), ("讲次", content.session_label), ("版本", "可编辑稿")]
    for index, (label, value) in enumerate(metadata):
        row, column = divmod(index, 2)
        _replace_cell(table.cell(row, column * 2), label)
        _replace_cell(table.cell(row, column * 2 + 1), value)
    sections = [
        ("一、教学目标", "\n".join(f"{i}. {item}" for i, item in enumerate(content.objectives, 1))),
        ("二、知识点", "\n".join(f"{i}. {item}" for i, item in enumerate(content.knowledge_points, 1))),
        ("三、教学重点", "\n".join(f"{i}. {item}" for i, item in enumerate(content.key_points, 1))),
        ("四、教学难点", "\n".join(f"{i}. {item}" for i, item in enumerate(content.difficult_points, 1))),
        ("五、教学方法与手段", f"教学方法：{'、'.join(content.methods)}\n教学手段：{'、'.join(content.tools)}"),
        ("六、课程思政融入", "；".join(content.ideological_elements)),
        ("七、教学过程", content.teaching_process),
        ("八、评价设计", content.assessment),
        ("九、教学后记", content.postscript),
    ]
    for heading, value in sections:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(heading), 12, True)
        for line in (value or "待教师完善").splitlines():
            paragraph = document.add_paragraph()
            _set_run_font(paragraph.add_run(line))
    paragraph = document.add_paragraph()
    _set_run_font(paragraph.add_run("附：数据来源"), 12, True)
    for index, reference in enumerate(references, 1):
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f"{index}. [{reference['layer']}] {reference['source_name']}  {reference.get('locator', '')}"), 9)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_docx(record: dict, template_path: Path | None = None, preserve_source_format: bool = True) -> tuple[bytes, str]:
    content = CourseDesignContent.model_validate(record["content"])
    if template_path and template_path.suffix.lower() == ".docx":
        rendered, report = _render_template(template_path, content)
        if report["compatible"]:
            return rendered, "source-template"
        if preserve_source_format:
            raise ValueError(report["message"] + "；请更换模板或关闭“保持原格式”后使用内置模板")
    return _standard_document(content, record.get("source_references", [])), "standard-template"
