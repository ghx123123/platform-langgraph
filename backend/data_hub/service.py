import hashlib
import html
import json
import os
import re
import shutil
from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from backend.documents.storage import original_path
from docx.oxml.ns import qn
from docx.shared import Pt

from backend.course_archives.models import ArchiveManifestItem
from backend.course_archives.service import (
    analyze_course_archive, normalize_archive_path,
)
from backend.documents.service import parse_document
from backend.documents.storage import persist_original


EXCLUDED_DIRECTORIES = {
    ".git", ".venv", "venv", "node_modules", "dist", "build", "__pycache__",
    ".cache", ".runtime", ".idea", ".vscode",
}
MAX_LOCAL_FILES = 6000


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def unit_id(archive_id: str, chapter: str | None) -> str:
    digest = hashlib.sha1((chapter or "全课程").encode("utf-8")).hexdigest()[:10]
    return f"{archive_id}:{digest}"


def rename_academic_term(archives: list[dict], current_name: str, new_name: str) -> list[dict]:
    source = current_name.strip() or "未设置学期"
    target = new_name.strip()
    if not target:
        raise ValueError("学期名称不能为空")
    stored_target = "" if target == "未设置学期" else target
    timestamp = utc_now()
    updated: list[dict] = []
    for archive in archives:
        label = str(archive.get("academic_term") or "").strip() or "未设置学期"
        if label != source:
            continue
        record = deepcopy(archive)
        record.update({"academic_term": stored_target, "updated_at": timestamp})
        updated.append(record)
    if not updated:
        raise KeyError("未找到需要重命名的学期目录")
    return updated


def normalize_library_root_name(name: str) -> str:
    normalized = re.sub(r"\s+", " ", name).strip()
    if not normalized:
        raise ValueError("资料库文件夹名称不能为空")
    if normalized in {".", ".."} or any(character in normalized for character in "/\\"):
        raise ValueError("资料库文件夹名称不能包含斜杠")
    return normalized[:100]


def create_library_root(archives: list[dict], name: str, document_store_path: Path) -> dict:
    normalized = normalize_library_root_name(name)
    if any(str(item.get("name") or "").strip().casefold() == normalized.casefold() for item in archives):
        raise FileExistsError("资料库根目录已存在同名文件夹")
    return analyze_course_archive(
        normalized,
        [],
        [],
        document_store_path,
        metadata={"course_title": normalized, "academic_term": "", "course_code": ""},
        extract_uploads=False,
    )


def rename_library_root(archives: list[dict], archive_id: str, name: str) -> dict:
    normalized = normalize_library_root_name(name)
    existing = next((item for item in archives if item.get("id") == archive_id), None)
    if not existing:
        raise KeyError("未找到资料库文件夹")
    if any(
        item.get("id") != archive_id
        and str(item.get("name") or "").strip().casefold() == normalized.casefold()
        for item in archives
    ):
        raise FileExistsError("资料库根目录已存在同名文件夹")
    updated = deepcopy(existing)
    updated.update({"name": normalized, "updated_at": utc_now()})
    return updated


def _content_preview(value: str, limit: int = 360) -> str:
    return re.sub(r"\s+", " ", value or "").strip()[:limit]


def _design_markdown(content: dict) -> str:
    sections = [
        ("教学目标", "\n".join(f"- {item}" for item in content.get("objectives", []))),
        ("教学重点", "\n".join(f"- {item}" for item in content.get("key_points", []))),
        ("教学难点", "\n".join(f"- {item}" for item in content.get("difficult_points", []))),
        ("教学方法", "、".join(content.get("methods", []))),
        ("教学手段", "、".join(content.get("tools", []))),
        ("教学过程", content.get("teaching_process", "")),
        ("评价设计", content.get("assessment", "")),
    ]
    return "\n\n".join(f"## {title}\n{value}" for title, value in sections if value)


def _folder_name(value: str) -> str:
    name = re.sub(r"\s+", " ", value or "").strip()
    if not name:
        raise ValueError("文件夹名称不能为空")
    if name in {".", ".."} or any(character in name for character in "/\\"):
        raise ValueError("文件夹名称不能包含斜杠")
    return name[:120]


def _folder(layout: dict, folder_id: str) -> dict:
    result = next((item for item in layout.get("folders", []) if item["id"] == folder_id), None)
    if not result:
        raise KeyError("未找到资料文件夹")
    return result


def _check_folder_destination(layout: dict, name: str, parent_id: str | None, system_parent: str | None, exclude_id: str | None = None) -> None:
    if parent_id and system_parent:
        raise ValueError("文件夹不能同时属于自定义目录和系统分类")
    if parent_id:
        _folder(layout, parent_id)
    duplicate = next((item for item in layout.get("folders", []) if item["id"] != exclude_id and item.get("parent_id") == parent_id and item.get("system_parent") == system_parent and item["name"].casefold() == name.casefold()), None)
    if duplicate:
        raise FileExistsError("目标位置已存在同名文件夹")


def create_data_folder(layout: dict, name: str, parent_id: str | None = None, system_parent: str | None = None) -> tuple[dict, dict]:
    result = deepcopy(layout)
    normalized = _folder_name(name)
    _check_folder_destination(result, normalized, parent_id, system_parent)
    timestamp = utc_now()
    folder = {"id": str(uuid4()), "unit_id": result["unit_id"], "name": normalized, "parent_id": parent_id, "system_parent": system_parent, "created_at": timestamp, "updated_at": timestamp}
    result.setdefault("folders", []).append(folder)
    result["updated_at"] = timestamp
    return result, folder


def update_data_folder(layout: dict, folder_id: str, name: str | None = None, move: bool = False, parent_id: str | None = None, system_parent: str | None = None) -> dict:
    result = deepcopy(layout)
    folder = _folder(result, folder_id)
    next_name = _folder_name(name) if name is not None else folder["name"]
    next_parent = parent_id if move else folder.get("parent_id")
    next_system = system_parent if move else folder.get("system_parent")
    if next_parent == folder_id:
        raise ValueError("文件夹不能移动到自身")
    ancestor = next_parent
    while ancestor:
        if ancestor == folder_id:
            raise ValueError("文件夹不能移动到自身的子文件夹")
        ancestor = _folder(result, ancestor).get("parent_id")
    _check_folder_destination(result, next_name, next_parent, next_system, folder_id)
    folder.update({"name": next_name, "parent_id": next_parent, "system_parent": next_system, "updated_at": utc_now()})
    result["updated_at"] = folder["updated_at"]
    return result


def delete_data_folder(layout: dict, folder_id: str) -> dict:
    result = deepcopy(layout)
    _folder(result, folder_id)
    if any(item.get("parent_id") == folder_id for item in result.get("folders", [])):
        raise RuntimeError("文件夹包含子文件夹，请先移动或删除子文件夹")
    if folder_id in result.get("placements", {}).values():
        raise RuntimeError("文件夹中仍有资料，请先移动资料")
    result["folders"] = [item for item in result.get("folders", []) if item["id"] != folder_id]
    result["updated_at"] = utc_now()
    return result


def folder_subtree_ids(layout: dict, folder_id: str) -> set[str]:
    _folder(layout, folder_id)
    result = {folder_id}
    changed = True
    while changed:
        changed = False
        for item in layout.get("folders", []):
            if item.get("parent_id") in result and item["id"] not in result:
                result.add(item["id"])
                changed = True
    return result


def delete_data_folder_recursive(layout: dict, folder_id: str) -> tuple[dict, set[str]]:
    result = deepcopy(layout)
    folder_ids = folder_subtree_ids(result, folder_id)
    block_ids = {block_id for block_id, placed_folder_id in result.get("placements", {}).items() if placed_folder_id in folder_ids}
    result["folders"] = [item for item in result.get("folders", []) if item["id"] not in folder_ids]
    result["placements"] = {block_id: value for block_id, value in result.get("placements", {}).items() if block_id not in block_ids}
    result["titles"] = {block_id: value for block_id, value in result.get("titles", {}).items() if block_id not in block_ids}
    result["updated_at"] = utc_now()
    return result, block_ids


def remove_blocks_from_layout(layout: dict, block_ids: set[str]) -> dict:
    result = deepcopy(layout)
    result["placements"] = {key: value for key, value in result.get("placements", {}).items() if key not in block_ids}
    result["titles"] = {key: value for key, value in result.get("titles", {}).items() if key not in block_ids}
    result["updated_at"] = utc_now()
    return result


def ensure_data_folder_path(
    layout: dict,
    parts: list[str],
    parent_id: str | None = None,
    system_parent: str | None = None,
) -> tuple[dict, str | None, int]:
    result = deepcopy(layout)
    current_parent = parent_id
    current_system = None if parent_id else system_parent
    created = 0
    for raw_name in parts:
        name = _folder_name(raw_name)
        existing = next((
            item for item in result.get("folders", [])
            if item.get("parent_id") == current_parent
            and item.get("system_parent") == current_system
            and item["name"].casefold() == name.casefold()
        ), None)
        if existing:
            current_parent = existing["id"]
        else:
            result, folder = create_data_folder(result, name, current_parent, current_system)
            current_parent = folder["id"]
            created += 1
        current_system = None
    return result, current_parent, created


def update_block_layout(layout: dict, block_id: str, title: str | None = None, move: bool = False, folder_id: str | None = None) -> dict:
    result = deepcopy(layout)
    if title is not None:
        normalized = re.sub(r"\s+", " ", title).strip()
        if not normalized:
            raise ValueError("资料名称不能为空")
        result.setdefault("titles", {})[block_id] = normalized[:240]
    if move:
        if folder_id:
            _folder(result, folder_id)
            result.setdefault("placements", {})[block_id] = folder_id
        else:
            result.setdefault("placements", {}).pop(block_id, None)
    result["updated_at"] = utc_now()
    return result


def apply_layouts(catalog: dict, layouts: list[dict]) -> dict:
    result = deepcopy(catalog)
    unit_ids = {item["id"] for item in result.get("units", [])}
    by_unit = {item["unit_id"]: item for item in layouts if item.get("unit_id") in unit_ids}
    folder_ids_by_unit = {
        unit_id: {item["id"] for item in layout.get("folders", [])}
        for unit_id, layout in by_unit.items()
    }
    folders: list[dict] = []
    for layout in by_unit.values():
        folders.extend(layout.get("folders", []))
    for block in result.get("blocks", []):
        layout = by_unit.get(block.get("unit_id"))
        if not layout:
            continue
        if block["id"] in layout.get("titles", {}):
            block["title"] = layout["titles"][block["id"]]
        folder_id = layout.get("placements", {}).get(block["id"])
        if folder_id and folder_id in folder_ids_by_unit.get(block.get("unit_id"), set()):
            block["folder_id"] = folder_id
    result["folders"] = folders
    return result


def organize_imported_archive(catalog: dict, layouts: list[dict], archive_id: str, folder_name: str) -> tuple[list[dict], dict]:
    normalized = _folder_name(folder_name)
    units = [item for item in catalog.get("units", []) if item.get("archive_id") == archive_id]
    if not units:
        raise KeyError("未找到刚导入的课程目录")
    by_unit = {item["unit_id"]: deepcopy(item) for item in layouts}
    changed: list[dict] = []
    folder_count = 0
    block_count = 0
    for unit in units:
        layout = by_unit.get(unit["id"], {"unit_id": unit["id"], "folders": [], "placements": {}, "titles": {}, "updated_at": ""})
        unit_blocks = [item for item in catalog.get("blocks", []) if item.get("unit_id") == unit["id"] and item.get("kind") == "original"]
        if unit_blocks:
            folder = next((item for item in layout.get("folders", []) if item.get("parent_id") is None and item.get("system_parent") is None and item["name"].casefold() == normalized.casefold()), None)
            if not folder:
                layout, folder = create_data_folder(layout, normalized)
                folder_count += 1
            for block in unit_blocks:
                layout = update_block_layout(layout, block["id"], move=True, folder_id=folder["id"])
                block_count += 1
        changed.append(layout)
    return changed, {"archive_id": archive_id, "unit_count": len(units), "folder_count": folder_count, "block_count": block_count}


def organize_source_archive(catalog: dict, layouts: list[dict], archive_id: str, source: dict) -> tuple[list[dict], dict]:
    units = [item for item in catalog.get("units", []) if item.get("archive_id") == archive_id]
    by_unit = {item["unit_id"]: deepcopy(item) for item in layouts}
    changed_layouts: list[dict] = []
    folder_count = 0
    block_count = 0
    for unit in units:
        layout = by_unit.get(unit["id"], {"unit_id": unit["id"], "folders": [], "placements": {}, "titles": {}, "updated_at": ""})
        mount_parent_id = source.get("mount_parent_id")
        if mount_parent_id and not any(item["id"] == mount_parent_id for item in layout.get("folders", [])):
            raise KeyError("导入位置已不存在，请返回目标文件夹后重新导入")
        existing_root = next((item for item in layout.get("folders", []) if item.get("source_folder_id") == source["id"]), None)
        unit_blocks = [
            item for item in catalog.get("blocks", [])
            if item.get("unit_id") == unit["id"]
            and item.get("kind") == "original"
            and item.get("source_folder_id") == source["id"]
        ]
        root = None
        target_root_id = mount_parent_id
        if source.get("selection_kind", "folder") == "folder":
            if existing_root:
                layout = update_data_folder(
                    layout, existing_root["id"], source["name"], True, mount_parent_id, None,
                )
                root = next(item for item in layout["folders"] if item["id"] == existing_root["id"])
            else:
                layout, root = create_data_folder(layout, source["name"], mount_parent_id)
                folder_count += 1
            root.update({
                "source_folder_id": source["id"], "source_kind": source["kind"],
                "source_selection_kind": source.get("selection_kind", "folder"),
                "source_path": source.get("root_path"), "last_scanned_at": source.get("last_scanned_at"),
            })
            target_root_id = root["id"]
        for directory_path in sorted(source.get("directory_paths", []), key=lambda value: (value.count("/"), value.casefold())):
            parts = [part for part in str(directory_path).replace("\\", "/").split("/") if part]
            if not parts:
                continue
            layout, _, created = ensure_data_folder_path(layout, parts, target_root_id, None)
            folder_count += created
        for block in unit_blocks:
            relative = str(block.get("source_relative_path") or block.get("locator") or block["title"])
            parts = [part for part in relative.replace("\\", "/").split("/")[:-1] if part]
            target_id = target_root_id
            if parts:
                layout, target_id, created = ensure_data_folder_path(layout, parts, target_root_id, None)
                folder_count += created
            layout = update_block_layout(layout, block["id"], move=True, folder_id=target_id)
            block_count += 1
        changed_layouts.append(layout)
    return changed_layouts, {
        "archive_id": archive_id, "source_id": source["id"],
        "unit_count": len(changed_layouts), "folder_count": folder_count, "block_count": block_count,
    }


def build_catalog(
    archives: list[dict],
    designs: list[dict],
    runs: list[dict],
    compositions: list[dict] | None = None,
    include_blocks: bool = True,
    target_unit_id: str | None = None,
) -> dict:
    designs_by_archive: dict[str, list[dict]] = {}
    run_design: dict[str, dict] = {}
    for design in designs:
        designs_by_archive.setdefault(design["archive_id"], []).append(design)
        if design.get("run_id"):
            run_design[design["run_id"]] = design
    run_context: dict[str, dict] = {}
    for run in runs:
        linked = run_design.get(run.get("id"))
        if linked:
            run_context[run["id"]] = {"archive_id": linked["archive_id"], "chapter": linked.get("chapter"), "design_id": linked["id"]}
            continue
        objective = run.get("objective", "")
        archive = next((item for item in archives if item.get("course_title") and item["course_title"] in objective), None)
        if not archive:
            continue
        chapter_match = re.search(r"第\s*(?:[0-9]{1,2}|[一二三四五六七八九十]+)\s*章", objective)
        chapter = chapter_match.group(0).replace(" ", "") if chapter_match else None
        matched_design = next((item for item in designs_by_archive.get(archive["id"], []) if not chapter or item.get("chapter") == chapter), None)
        run_context[run["id"]] = {"archive_id": archive["id"], "chapter": chapter, "design_id": matched_design.get("id") if matched_design else None}

    blocks: list[dict] = []
    units: list[dict] = []
    for archive in archives:
        materials = archive.get("materials", [])
        sources_by_id = {item["id"]: item for item in archive.get("source_folders", [])}
        archive_designs = designs_by_archive.get(archive["id"], [])
        course_unit_id = unit_id(archive["id"], None)
        linked_runs = {
            run_id for run_id, context in run_context.items()
            if context["archive_id"] == archive["id"]
        }
        units.append({
            "id": course_unit_id,
            "archive_id": archive["id"],
            "archive_name": archive.get("name") or archive.get("course_title") or "未命名资料库",
            "academic_term": archive.get("academic_term") or "未设置学期",
            "course_title": archive.get("course_title", archive.get("name", "未命名课程")),
            "course_code": archive.get("course_code", ""),
            "chapter": "全课程",
            "material_count": len(materials),
            "parsed_count": sum(1 for item in materials if item.get("parse_status") == "parsed"),
            "design_count": len(archive_designs),
            "generated_count": len(linked_runs),
            "updated_at": max([archive.get("updated_at", ""), *[item.get("updated_at", "") for item in archive_designs]]),
        })
        if not include_blocks:
            continue
        for material in materials:
            if target_unit_id and course_unit_id != target_unit_id:
                continue
            base = {
                "archive_id": archive["id"], "unit_id": course_unit_id, "source_name": material["name"],
                "updated_at": archive.get("updated_at", ""), "editable": False,
                "category": material.get("category", "other"),
                "modified_at": material.get("last_modified"),
                "source_folder_id": material.get("source_folder_id"),
                "source_kind": material.get("source_kind"),
                "source_selection_kind": (sources_by_id.get(material.get("source_folder_id")) or {}).get("selection_kind"),
                "source_relative_path": material.get("source_relative_path"),
            }
            blocks.append({
                **base, "id": f"material:{archive['id']}:{material['id']}:original", "kind": "original",
                "title": material["name"], "content_preview": material.get("path", ""),
                "locator": material.get("path", ""),
                "original_url": f"/api/documents/{material['document_id']}/original" if material.get("document_id") else None,
                "preview_url": f"/api/documents/{material['document_id']}/preview" if material.get("preview_available") else None,
            })
            if material.get("parse_status") == "parsed":
                blocks.append({
                    **base, "id": f"material:{archive['id']}:{material['id']}:extracted", "kind": "extracted",
                    "title": f"{material['name']} · 提取正文", "content_preview": material.get("excerpt", ""),
                    "locator": f"material:{material['id']}:extracted",
                    "original_url": f"/api/documents/{material['document_id']}/original" if material.get("document_id") else None,
                    "preview_url": f"/api/documents/{material['document_id']}/preview" if material.get("preview_available") else None,
                })

    for design in designs if include_blocks else []:
        content = design.get("content", {})
        uid = unit_id(design["archive_id"], None)
        if target_unit_id and uid != target_unit_id:
            continue
        markdown = _design_markdown(content)
        blocks.append({
            "id": f"design:{design['id']}:teaching", "kind": "teaching_design", "title": f"{design['title']} · 教学设计",
            "content_preview": _content_preview(markdown), "content": markdown, "archive_id": design["archive_id"],
            "unit_id": uid, "design_id": design["id"], "run_id": design.get("run_id"), "source_name": design["title"],
            "locator": f"course-design:{design['id']}:v{design.get('version', 1)}", "editable": True,
            "updated_at": design.get("updated_at", ""),
        })
        ideology = "\n".join(f"- {item}" for item in content.get("ideological_elements", []))
        if ideology:
            blocks.append({
                "id": f"design:{design['id']}:ideology", "kind": "ideological_element", "title": f"{design['title']} · 课程思政元素",
                "content_preview": _content_preview(ideology), "content": ideology, "archive_id": design["archive_id"],
                "unit_id": uid, "design_id": design["id"], "run_id": design.get("run_id"), "source_name": design["title"],
                "locator": f"course-design:{design['id']}:ideological-elements", "editable": True,
                "updated_at": design.get("updated_at", ""),
            })

    kind_by_phase = {
        "student_question": "student_question", "teacher_answer": "teacher_answer", "supervisor_comment": "supervisor_review",
    }
    for run in runs if include_blocks else []:
        context = run_context.get(run.get("id"))
        if not context:
            continue
        uid = unit_id(context["archive_id"], None)
        if target_unit_id and uid != target_unit_id:
            continue
        framework = run.get("teaching_data", {}).get("teaching_framework", {}) or {}
        if framework:
            framework_content = "\n\n".join(filter(None, (
                "## 学习目标\n" + "\n".join(f"- {item}" for item in framework.get("learning_objectives", [])),
                "## 教学策略\n" + "\n".join(f"- {item}" for item in framework.get("strategies", [])),
                "## 教学环节\n" + "\n".join(f"- {item.get('name', '教学环节')}：{item.get('purpose', '')} {item.get('activity', '')}" for item in framework.get("stages", [])),
            )))
            blocks.append({
                "id": f"run:{run['id']}:teaching_design:framework", "kind": "teaching_design", "title": f"{run.get('objective', '教学会话')} · 生成教学设计",
                "content_preview": _content_preview(framework_content), "content": framework_content,
                "archive_id": context["archive_id"], "unit_id": uid, "design_id": context.get("design_id"), "run_id": run["id"],
                "source_name": run.get("objective", "教学会话"), "locator": f"workflow-run:{run['id']}:teaching-framework",
                "editable": True, "updated_at": str(run.get("updated_at", "")),
            })
        for index, message in enumerate(run.get("teaching_data", {}).get("messages", [])):
            kind = kind_by_phase.get(message.get("phase"))
            if not kind:
                continue
            title = f"{message.get('agent_name', message.get('agent_type', '智能体'))} · 第 {message.get('iteration', 0)} 轮"
            content = message.get("content", "")
            blocks.append({
                "id": f"run:{run['id']}:{kind}:{message.get('id', index)}", "kind": kind, "title": title,
                "content_preview": _content_preview(content), "content": content, "archive_id": context["archive_id"],
                "unit_id": uid, "design_id": context.get("design_id"), "run_id": run["id"], "source_name": run.get("objective", "教学会话"),
                "locator": f"workflow-run:{run['id']}:{message.get('phase')}:{index}", "editable": True,
                "updated_at": str(run.get("updated_at", "")),
            })

    for composition in (compositions or []) if include_blocks else []:
        composition_unit_id = unit_id(composition["archive_id"], None) if composition.get("archive_id") else composition.get("unit_id")
        if target_unit_id and composition_unit_id != target_unit_id:
            continue
        for item in composition.get("blocks", []):
            blocks.append({
                "id": f"composition:{composition['id']}:{item['id']}", "kind": item.get("kind", "imported"),
                "title": item.get("title", composition["title"]), "content_preview": _content_preview(item.get("content", "")),
                "content": item.get("content", ""), "archive_id": composition.get("archive_id"),
                "unit_id": composition_unit_id, "source_name": composition["title"],
                "locator": f"composition:{composition['id']}:v{composition.get('version', 1)}",
                "editable": True, "updated_at": composition.get("updated_at", ""),
            })

    generated = sum(1 for item in blocks if item["kind"] not in {"original", "extracted"})
    return {
        "stats": {
            "terms": len({item["academic_term"] for item in units}),
            "courses": len({item["archive_id"] for item in units}),
            "units": len(units), "materials": sum(len(item.get("materials", [])) for item in archives),
            "generated_blocks": generated,
        },
        "terms": sorted({item["academic_term"] for item in units}),
        "courses": sorted({item["course_title"] for item in units}),
        "units": sorted(units, key=lambda item: (item["academic_term"], item["course_title"], item["chapter"])),
        "blocks": sorted(blocks, key=lambda item: item.get("updated_at", ""), reverse=True),
    }


def filter_catalog(catalog: dict, query: str = "", term: str = "", course: str = "", kind: str = "") -> dict:
    query = query.strip().lower()
    archive_ids = {
        item["archive_id"] for item in catalog["units"]
        if (not term or item["academic_term"] == term) and (not course or item["course_title"] == course)
    }
    units = [item for item in catalog["units"] if item["archive_id"] in archive_ids]
    unit_ids = {item["id"] for item in units}
    folders = [item for item in catalog.get("folders", []) if item.get("unit_id") in unit_ids]
    blocks = [item for item in catalog["blocks"] if item.get("unit_id") in unit_ids and (not kind or item["kind"] == kind)]
    if query:
        matching_blocks = [item for item in blocks if query in " ".join((item["title"], item.get("source_name", ""), item.get("content_preview", ""), item.get("locator", ""))).lower()]
        matching_unit_ids = {item.get("unit_id") for item in matching_blocks}
        units = [
            item for item in units
            if item["id"] in matching_unit_ids
            or query in " ".join((item["academic_term"], item["course_title"], item["course_code"], item["chapter"])).lower()
        ]
        blocks = matching_blocks
    return {**catalog, "units": units, "folders": folders, "blocks": blocks}


def block_detail(block_id: str, catalog: dict, archives: list[dict]) -> dict:
    block = next((item for item in catalog["blocks"] if item["id"] == block_id), None)
    if not block:
        raise KeyError("未找到内容块")
    result = deepcopy(block)
    if result["kind"] == "extracted":
        _, archive_id, material_id, _ = block_id.split(":", 3)
        archive = next(item for item in archives if item["id"] == archive_id)
        result["content"] = archive.get("_documents", {}).get(material_id, {}).get("raw_text", "")
        # S5: 原文(解析后文本)预览 — 供前端"原文预览"查看
        result["text"] = result["content"][:8000] if result["content"] else ""
        result["parse_status"] = (archive.get("_documents", {}).get(material_id, {}) or {}).get("parse_status", "metadata_only")
    elif result["kind"] == "original":
        result["content"] = result.get("locator", "")
        # S5: original 也可能已有解析结果 (_documents[material_id]) — 补 text/parse_status
        _, archive_id, material_id, _ = block_id.split(":", 3)
        archive = next((item for item in archives if item["id"] == archive_id), None)
        if archive:
            doc = archive.get("_documents", {}).get(material_id, {})
            raw = doc.get("raw_text", "")
            result["text"] = raw[:8000] if raw else ""
            result["parse_status"] = doc.get("parse_status") or None
            if raw:
                result["content"] = raw
    return result


def register_source_manifest(
    existing: dict | None,
    source_name: str,
    source_kind: str,
    manifest: list[ArchiveManifestItem],
    metadata: dict,
    document_store: Path,
    source_id: str | None = None,
    root_path: str | None = None,
    selection_kind: str = "folder",
    directory_paths: list[str] | None = None,
    mount_parent_id: str | None = None,
) -> tuple[dict, dict, dict, list[str]]:
    timestamp = utc_now()
    sources = [dict(item) for item in (existing or {}).get("source_folders", [])]
    source = next((item for item in sources if item.get("id") == source_id), None)
    if source_id and not source:
        raise KeyError("未找到需要刷新的来源文件夹")
    if source is None and source_kind == "local" and root_path:
        resolved = str(Path(root_path).resolve()).casefold()
        source = next((item for item in sources if item.get("kind") == "local" and str(item.get("root_path") or "").casefold() == resolved), None)
    normalized_name = _folder_name(source_name or (Path(root_path).name if root_path else "导入资料"))
    if source is None:
        existing_names = {item.get("name", "").casefold() for item in sources}
        base_name = normalized_name
        suffix = 2
        while normalized_name.casefold() in existing_names:
            normalized_name = f"{base_name} ({suffix})"
            suffix += 1
        source = {
            "id": str(uuid4()), "name": normalized_name[:160], "kind": source_kind,
            "selection_kind": selection_kind,
            "root_path": root_path, "file_count": 0, "directory_count": 0, "directory_paths": [],
            "mount_parent_id": mount_parent_id, "last_scanned_at": None,
            "created_at": timestamp, "updated_at": timestamp,
        }
        sources.append(source)
    source_id = source["id"]
    source_name = source["name"]

    normalized_directories: list[str] = []
    seen_directories: set[str] = set()
    for raw_path in directory_paths or []:
        relative = normalize_archive_path(raw_path)
        if not relative:
            continue
        if relative.casefold().startswith(f"{source_name}/".casefold()):
            relative = relative[len(source_name) + 1:]
        key = relative.casefold()
        if key in seen_directories:
            continue
        seen_directories.add(key)
        normalized_directories.append(relative)

    normalized_manifest: list[ArchiveManifestItem] = []
    for item in manifest:
        relative = normalize_archive_path(item.path)
        if not relative:
            continue
        if relative.casefold().startswith(f"{source_name}/".casefold()):
            relative = relative[len(source_name) + 1:]
        normalized_manifest.append(ArchiveManifestItem(path=relative, size=item.size, last_modified=item.last_modified))

    previous_source = [item for item in (existing or {}).get("materials", []) if item.get("source_folder_id") == source_id]
    if existing is not None and not normalized_manifest and not previous_source:
        source.update({
            "kind": source_kind,
            "selection_kind": selection_kind,
            "root_path": root_path if source_kind == "local" else None,
            "file_count": 0,
            "directory_count": len(normalized_directories),
            "directory_paths": normalized_directories,
            "mount_parent_id": mount_parent_id if mount_parent_id is not None else source.get("mount_parent_id"),
            "last_scanned_at": timestamp,
            "updated_at": timestamp,
        })
        record = deepcopy(existing)
        record["source_folders"] = sources
        record["updated_at"] = timestamp
        if source_kind == "local":
            record["local_root"] = root_path
            record["last_scanned_at"] = timestamp
        return record, source, {
            "added": 0, "changed": 0, "unchanged": 0, "removed": 0,
            "parsed": record.get("parsed_files", 0), "_changed_paths": set(),
        }, []
    current_full_paths = {f"{source_name}/{item.path}".casefold() for item in normalized_manifest}
    if not previous_source:
        previous_source = [
            item for item in (existing or {}).get("materials", [])
            if not item.get("source_folder_id") and item.get("path", "").casefold() in current_full_paths
        ]
    previous_by_relative = {
        (item.get("source_relative_path") or item.get("path", "")[len(source_name) + 1:]).casefold(): item
        for item in previous_source
    }
    current_by_relative = {item.path.casefold(): item for item in normalized_manifest}
    added = set(current_by_relative) - set(previous_by_relative)
    removed = set(previous_by_relative) - set(current_by_relative)
    changed = {
        path for path in set(current_by_relative) & set(previous_by_relative)
        if (current_by_relative[path].size, current_by_relative[path].last_modified)
        != (previous_by_relative[path].get("size", 0), previous_by_relative[path].get("last_modified"))
    }
    relocated_previous: dict[str, dict] = {}
    if source_kind == "local" and added and removed:
        added_by_signature: dict[tuple[str, int, int | str | None], list[str]] = {}
        removed_by_signature: dict[tuple[str, int, int | str | None], list[str]] = {}
        for path in added:
            item = current_by_relative[path]
            signature = (Path(item.path).name.casefold(), item.size, item.last_modified)
            added_by_signature.setdefault(signature, []).append(path)
        for path in removed:
            item = previous_by_relative[path]
            signature = (Path(item.get("source_relative_path") or item.get("path", "")).name.casefold(), item.get("size", 0), item.get("last_modified"))
            removed_by_signature.setdefault(signature, []).append(path)
        relocated_old_paths: set[str] = set()
        for signature, new_paths in added_by_signature.items():
            old_paths = removed_by_signature.get(signature, [])
            if len(new_paths) == 1 and len(old_paths) == 1:
                relocated_previous[new_paths[0]] = previous_by_relative[old_paths[0]]
                relocated_old_paths.add(old_paths[0])
        added -= set(relocated_previous)
        removed -= relocated_old_paths
        changed |= set(relocated_previous)

    source_material_ids = {item.get("id") for item in previous_source}
    preserved = [item for item in (existing or {}).get("materials", []) if item.get("id") not in source_material_ids]
    combined_manifest = [
        ArchiveManifestItem(path=item["path"], size=item.get("size", 0), last_modified=item.get("last_modified"))
        for item in preserved
    ]
    combined_manifest.extend(
        ArchiveManifestItem(path=f"{source_name}/{item.path}", size=item.size, last_modified=item.last_modified)
        for item in normalized_manifest
    )
    archive_name = (existing or {}).get("name") or metadata.get("archive_name") or metadata.get("course_title") or source_name
    record = analyze_course_archive(
        archive_name, combined_manifest, [], document_store, existing, metadata, extract_uploads=False,
    )
    preserved_by_path = {item.get("path", "").casefold(): item for item in preserved}
    id_replacements: dict[str, str] = {}
    for material in record.get("materials", []):
        relative = None
        prefix = f"{source_name}/"
        if material.get("path", "").casefold().startswith(prefix.casefold()):
            relative = material["path"][len(prefix):]
        if relative is not None and relative.casefold() in current_by_relative:
            relative_key = relative.casefold()
            previous = previous_by_relative.get(relative_key) or relocated_previous.get(relative_key)
            if previous and relative_key in relocated_previous and material["id"] != previous["id"]:
                id_replacements[material["id"]] = previous["id"]
                material["id"] = previous["id"]
            if previous and previous.get("document_id") and not material.get("document_id"):
                for key in ("sha256", "parse_status", "parse_message", "document_id", "preview_available", "character_count", "excerpt"):
                    material[key] = previous.get(key, material.get(key))
                previous_documents = (existing or {}).get("_documents", {})
                if previous["id"] in previous_documents:
                    record.setdefault("_documents", {})[material["id"]] = previous_documents[previous["id"]]
            material.update({
                "source_folder_id": source_id,
                "source_kind": source_kind,
                "source_relative_path": relative,
            })
        else:
            previous = preserved_by_path.get(material.get("path", "").casefold())
            if previous:
                for key in ("source_folder_id", "source_kind", "source_relative_path"):
                    if previous.get(key) is not None:
                        material[key] = previous[key]

    if id_replacements:
        def replace_material_references(value):
            if isinstance(value, str):
                return id_replacements.get(value, value)
            if isinstance(value, list):
                return [replace_material_references(item) for item in value]
            if isinstance(value, dict):
                return {
                    id_replacements.get(key, key): replace_material_references(item)
                    for key, item in value.items()
                }
            return value

        record = replace_material_references(record)

    source.update({
        "kind": source_kind,
        "selection_kind": selection_kind,
        "root_path": root_path if source_kind == "local" else None,
        "file_count": len(normalized_manifest),
        "directory_count": len(normalized_directories),
        "directory_paths": normalized_directories,
        "mount_parent_id": mount_parent_id if mount_parent_id is not None else source.get("mount_parent_id"),
        "last_scanned_at": timestamp,
        "updated_at": timestamp,
    })
    record["source_folders"] = sources
    if source_kind == "local":
        record["local_root"] = root_path
        record["last_scanned_at"] = timestamp
    removed_document_ids = [previous_by_relative[path].get("document_id") for path in removed if previous_by_relative[path].get("document_id")]
    changes = {
        "added": len(added), "changed": len(changed),
        "unchanged": len(current_by_relative) - len(added) - len(changed),
        "removed": len(removed), "parsed": record.get("parsed_files", 0),
        "_changed_paths": added | changed,
    }
    return record, source, changes, removed_document_ids


def scan_local_source(root_path: str, existing: dict | None, metadata: dict, document_store: Path, source_id: str | None = None, source_name: str = "") -> tuple[dict, dict, dict, list[str]]:
    root = Path(root_path).expanduser().resolve()
    if not root.is_dir():
        raise ValueError("本地资料路径不存在或不是文件夹")
    scan_warnings: list[str] = []

    def record_walk_error(error: OSError) -> None:
        if len(scan_warnings) < 20:
            scan_warnings.append(f"已跳过无法访问的目录：{error.filename or '未知路径'}")

    manifest: list[ArchiveManifestItem] = []
    directory_paths: list[str] = []
    for current_root, directory_names, file_names in os.walk(root, topdown=True, onerror=record_walk_error, followlinks=False):
        current = Path(current_root)
        available_directories: list[str] = []
        for name in directory_names:
            candidate = current / name
            if name.lower() in EXCLUDED_DIRECTORIES:
                continue
            try:
                if candidate.is_symlink():
                    continue
            except OSError:
                if len(scan_warnings) < 20:
                    scan_warnings.append(f"已跳过无法访问的目录：{candidate}")
                continue
            available_directories.append(name)
        directory_names[:] = available_directories
        for name in available_directories:
            candidate = current / name
            try:
                directory_paths.append(candidate.relative_to(root).as_posix())
            except ValueError:
                continue
        for name in file_names:
            item = current / name
            try:
                if item.is_symlink():
                    continue
                stat = item.stat()
                relative = item.relative_to(root).as_posix()
            except (OSError, ValueError) as exc:
                if len(scan_warnings) < 20:
                    scan_warnings.append(f"已跳过无法读取的文件：{item}（{exc}）")
                continue
            manifest.append(ArchiveManifestItem(path=relative, size=stat.st_size, last_modified=int(stat.st_mtime * 1000)))
            if len(manifest) > MAX_LOCAL_FILES:
                raise ValueError(f"单个本地资料源最多索引 {MAX_LOCAL_FILES} 个文件")
    record, source, changes, removed_document_ids = register_source_manifest(
        existing, source_name or root.name, "local", manifest,
        {**metadata, "local_root": str(root), "last_scanned_at": utc_now()},
        document_store, source_id, str(root), "folder",
        directory_paths,
    )
    if scan_warnings:
        record["warnings"] = [*record.get("warnings", []), *scan_warnings]
    changes.pop("_changed_paths")
    changes["parsed"] = 0
    return record, source, changes, removed_document_ids


def _local_source(archive: dict, source_id: str) -> tuple[dict, Path]:
    source = next((item for item in archive.get("source_folders", []) if item.get("id") == source_id), None)
    if not source or source.get("kind") != "local" or not source.get("root_path"):
        raise KeyError("未找到有效的本地来源")
    root = Path(source["root_path"]).expanduser().resolve()
    if not root.is_dir():
        raise FileNotFoundError("本地来源目录已不存在")
    return source, root


def record_local_deletions(archive: dict, material_ids: set[str], folder_paths: list[str] | None = None) -> dict:
    result = deepcopy(archive)
    source_ids = {item["id"] for item in result.get("source_folders", []) if item.get("kind") == "local"}
    deleted = result.setdefault("_local_sync_deletions", {})
    for material in result.get("materials", []):
        source_id = material.get("source_folder_id")
        relative = normalize_archive_path(str(material.get("source_relative_path") or ""))
        if material.get("id") in material_ids and source_id in source_ids and relative:
            deleted[f"{source_id}:file:{relative.casefold()}"] = {
                "source_id": source_id, "path": relative, "kind": "file",
                "size": material.get("size", 0), "recorded_at": utc_now(),
            }
    for entry in folder_paths or []:
        try:
            source_id, relative = entry.split("\u0000", 1)
        except ValueError:
            continue
        relative = normalize_archive_path(relative)
        if source_id in source_ids and relative:
            deleted[f"{source_id}:directory:{relative.casefold()}"] = {
                "source_id": source_id, "path": relative, "kind": "directory", "recorded_at": utc_now(),
            }
    return result


def local_source_diff(archive: dict, layout: dict, source_id: str, document_store: Path | None = None) -> dict:
    source, root = _local_source(archive, source_id)
    local_files: dict[str, dict] = {}
    local_directories: dict[str, str] = {}
    for current_root, directory_names, file_names in os.walk(root, topdown=True, followlinks=False):
        current = Path(current_root)
        directory_names[:] = [name for name in directory_names if name.lower() not in EXCLUDED_DIRECTORIES and not (current / name).is_symlink()]
        for name in directory_names:
            relative = (current / name).relative_to(root).as_posix()
            local_directories[relative.casefold()] = relative
        for name in file_names:
            path = current / name
            if path.is_symlink():
                continue
            stat = path.stat()
            relative = path.relative_to(root).as_posix()
            local_files[relative.casefold()] = {"path": relative, "size": stat.st_size, "modified": int(stat.st_mtime * 1000)}

    platform_files = {
        normalize_archive_path(str(item.get("source_relative_path") or "")).casefold(): item
        for item in archive.get("materials", [])
        if item.get("source_folder_id") == source_id and item.get("source_kind") == "local" and item.get("source_relative_path")
    }
    source_root = next((item for item in layout.get("folders", []) if item.get("source_folder_id") == source_id), None)
    folders = {item["id"]: item for item in layout.get("folders", [])}
    platform_directories: dict[str, str] = {}
    if source_root:
        for folder in layout.get("folders", []):
            if folder["id"] == source_root["id"]:
                continue
            parts: list[str] = []
            current = folder
            visited: set[str] = set()
            while current and current["id"] != source_root["id"] and current["id"] not in visited:
                visited.add(current["id"])
                parts.append(current["name"])
                current = folders.get(current.get("parent_id")) if current.get("parent_id") else None
            if current and current["id"] == source_root["id"]:
                relative = "/".join(reversed(parts))
                platform_directories[relative.casefold()] = relative

    tombstones = [item for item in archive.get("_local_sync_deletions", {}).values() if item.get("source_id") == source_id]
    deleted_files = {item["path"].casefold() for item in tombstones if item.get("kind") == "file"}
    deleted_directories = {item["path"].casefold() for item in tombstones if item.get("kind") == "directory"}
    items: list[dict] = []
    for key, local in local_files.items():
        platform = platform_files.get(key)
        if key in deleted_files:
            items.append({"path": local["path"], "kind": "file", "status": "platform_deleted", "local_size": local["size"], "can_restore": False})
        elif not platform:
            items.append({"path": local["path"], "kind": "file", "status": "local_added", "local_size": local["size"], "can_restore": False})
        elif platform.get("size") != local["size"] or platform.get("last_modified") != local["modified"]:
            document_id = platform.get("document_id")
            can_restore = bool(document_store and document_id and original_path(document_store, document_id).exists())
            items.append({"path": local["path"], "kind": "file", "status": "local_changed", "local_size": local["size"], "platform_size": platform.get("size"), "can_restore": can_restore})
    for key, platform in platform_files.items():
        if key not in local_files:
            document_id = platform.get("document_id")
            can_restore = bool(document_store and document_id and original_path(document_store, document_id).exists())
            items.append({"path": platform.get("source_relative_path") or platform["name"], "kind": "file", "status": "local_removed", "platform_size": platform.get("size"), "can_restore": can_restore})
    for key, relative in local_directories.items():
        if key in deleted_directories:
            items.append({"path": relative, "kind": "directory", "status": "platform_deleted", "can_restore": False})
        elif key not in platform_directories:
            items.append({"path": relative, "kind": "directory", "status": "local_directory_added", "can_restore": False})
    for key, relative in platform_directories.items():
        if key not in local_directories and key not in deleted_directories:
            items.append({"path": relative, "kind": "directory", "status": "platform_directory_added", "can_restore": True})
    order = {"platform_deleted": 0, "local_removed": 1, "local_changed": 2, "local_added": 3, "platform_directory_added": 4, "local_directory_added": 5}
    items.sort(key=lambda item: (order[item["status"]], item["path"].casefold()))
    return {
        "archive_id": archive["id"], "source_id": source_id, "source_name": source["name"], "local_root": str(root),
        "items": items,
        "local_changes": sum(item["status"].startswith("local_") for item in items),
        "platform_changes": sum(item["status"] in {"platform_deleted", "platform_directory_added"} for item in items),
        "blocked_restores": sum(item["status"] in {"local_removed", "local_changed"} and not item["can_restore"] for item in items),
        "checked_at": utc_now(),
    }


def apply_platform_to_local(archive: dict, layout: dict, source_id: str, document_store: Path) -> tuple[int, int]:
    source, root = _local_source(archive, source_id)
    diff = local_source_diff(archive, layout, source_id, document_store)
    applied = 0
    skipped = 0
    material_by_path = {
        normalize_archive_path(str(item.get("source_relative_path") or "")).casefold(): item
        for item in archive.get("materials", []) if item.get("source_folder_id") == source_id
    }
    tombstones = [item for item in archive.get("_local_sync_deletions", {}).values() if item.get("source_id") == source_id]
    deleted_directory_paths = [Path(item["path"]) for item in tombstones if item.get("kind") == "directory"]
    protected_by_deleted_directory = lambda value: any(Path(value) == directory or directory in Path(value).parents for directory in deleted_directory_paths)
    ordered_tombstones = sorted(tombstones, key=lambda item: (item.get("kind") == "directory", -len(Path(item["path"]).parts)))
    for item in ordered_tombstones:
        target = (root / item["path"]).resolve()
        target.relative_to(root)
        if item.get("kind") == "file":
            if target.is_file():
                target.unlink(); applied += 1
            elif target.exists():
                skipped += 1
        elif target.is_dir():
            try:
                target.rmdir(); applied += 1
            except OSError:
                skipped += 1
    for item in diff["items"]:
        target = (root / item["path"]).resolve()
        target.relative_to(root)
        if item["status"] == "platform_directory_added":
            if not target.exists():
                target.mkdir(parents=True); applied += 1
        elif item["status"] in {"local_removed", "local_changed"}:
            material = material_by_path.get(item["path"].casefold())
            document_id = material.get("document_id") if material else None
            stored = original_path(document_store, document_id) if document_id else None
            if stored and stored.is_file():
                target.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(stored, target); applied += 1
            else:
                skipped += 1
        elif item["status"] == "local_added" and target.is_file():
            if protected_by_deleted_directory(item["path"]):
                skipped += 1
            else:
                target.unlink(); applied += 1
    local_only_directories = sorted(
        (item for item in diff["items"] if item["status"] == "local_directory_added"),
        key=lambda item: len(Path(item["path"]).parts), reverse=True,
    )
    for item in local_only_directories:
        if protected_by_deleted_directory(item["path"]):
            continue
        target = (root / item["path"]).resolve()
        target.relative_to(root)
        if target.is_dir():
            try:
                target.rmdir(); applied += 1
            except OSError:
                skipped += 1
    source["last_scanned_at"] = utc_now()
    archive["_local_sync_deletions"] = {
        key: value for key, value in archive.get("_local_sync_deletions", {}).items() if value.get("source_id") != source_id
    }
    return applied, skipped


def transfer_local_materials(
    archive: dict,
    layout: dict,
    block_ids: list[str],
    destination_folder_id: str,
    operation: str,
) -> tuple[dict, int]:
    if operation not in {"copy", "move"}:
        raise ValueError("本机同步方式无效")
    folders = {item["id"]: item for item in layout.get("folders", [])}
    destination = folders.get(destination_folder_id)
    if not destination:
        raise KeyError("目标文件夹已不存在")

    relative_parts: list[str] = []
    current = destination
    visited: set[str] = set()
    source_root = None
    while current:
        if current["id"] in visited:
            raise ValueError("目标目录关系异常，请刷新后重试")
        visited.add(current["id"])
        if current.get("source_folder_id"):
            source_root = current
            break
        relative_parts.append(current["name"])
        current = folders.get(current.get("parent_id")) if current.get("parent_id") else None
    if not source_root or source_root.get("source_kind") != "local":
        raise ValueError("目标文件夹未关联本机目录，只能调整平台目录")

    source = next((item for item in archive.get("source_folders", []) if item.get("id") == source_root.get("source_folder_id")), None)
    if not source or source.get("kind") != "local" or not source.get("root_path"):
        raise ValueError("目标文件夹的本机来源已失效，请先刷新来源")
    root = Path(source["root_path"]).expanduser().resolve()
    if not root.is_dir():
        raise FileNotFoundError("本机来源目录已不存在")
    target_directory = root.joinpath(*reversed(relative_parts)).resolve()
    try:
        target_directory.relative_to(root)
    except ValueError as exc:
        raise ValueError("目标目录超出本机来源范围") from exc

    materials = {item["id"]: item for item in archive.get("materials", [])}
    planned: list[tuple[Path, Path]] = []
    destinations: set[str] = set()
    for block_id in block_ids:
        parts = block_id.split(":")
        if len(parts) != 4 or parts[0] != "material" or parts[1] != archive.get("id") or parts[3] != "original":
            raise ValueError("本机同步只支持当前资料库中的原始文件")
        material = materials.get(parts[2])
        if not material or material.get("source_folder_id") != source["id"] or material.get("source_kind") != "local":
            raise ValueError("所选文件必须全部来自目标文件夹关联的同一本机目录")
        relative = Path(str(material.get("source_relative_path") or material["name"]))
        source_path = (root / relative).resolve()
        try:
            source_path.relative_to(root)
        except ValueError as exc:
            raise ValueError("源文件路径超出本机来源范围") from exc
        if not source_path.is_file():
            raise FileNotFoundError(f"本机源文件已不存在：{material['name']}")
        destination_path = (target_directory / source_path.name).resolve()
        key = str(destination_path).casefold()
        if key in destinations:
            raise FileExistsError(f"多个文件在目标目录中重名：{source_path.name}")
        destinations.add(key)
        if source_path == destination_path:
            raise ValueError(f"“{source_path.name}”已经位于目标本机目录")
        if destination_path.exists():
            raise FileExistsError(f"本机目标目录已存在同名文件：{source_path.name}")
        planned.append((source_path, destination_path))

    target_directory.mkdir(parents=True, exist_ok=True)
    completed: list[tuple[Path, Path]] = []
    try:
        for source_path, destination_path in planned:
            if operation == "copy":
                shutil.copy2(source_path, destination_path)
            else:
                shutil.move(str(source_path), str(destination_path))
            completed.append((source_path, destination_path))
    except OSError:
        for source_path, destination_path in reversed(completed):
            try:
                if operation == "copy":
                    destination_path.unlink(missing_ok=True)
                elif destination_path.exists() and not source_path.exists():
                    shutil.move(str(destination_path), str(source_path))
            except OSError:
                pass
        raise
    return source, len(completed)


def resolve_local_folder_path(archive: dict, layout: dict, folder_id: str) -> tuple[dict, Path, str]:
    folders = {item["id"]: item for item in layout.get("folders", [])}
    current = folders.get(folder_id)
    if not current:
        raise KeyError("目标文件夹已不存在")
    parts: list[str] = []
    visited: set[str] = set()
    local_root = None
    while current:
        if current["id"] in visited:
            raise ValueError("目录关系异常，请刷新后重试")
        visited.add(current["id"])
        if current.get("source_folder_id") and current.get("source_kind") == "local":
            local_root = current
            break
        parts.append(current["name"])
        current = folders.get(current.get("parent_id")) if current.get("parent_id") else None
    if not local_root:
        raise ValueError("当前位置未关联本地来源，无法同步到本机")
    source = next((item for item in archive.get("source_folders", []) if item.get("id") == local_root.get("source_folder_id")), None)
    if not source or source.get("kind") != "local" or not source.get("root_path"):
        raise ValueError("本地来源已失效，请先刷新来源")
    root = Path(source["root_path"]).expanduser().resolve()
    if not root.is_dir():
        raise FileNotFoundError("本地来源目录已不存在")
    relative = "/".join(reversed(parts))
    target = root.joinpath(*reversed(parts)).resolve()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise ValueError("同步目标超出本地来源范围") from exc
    return source, target, relative


def sync_folder_to_local(archive: dict, layout: dict, folder_id: str) -> tuple[dict, int]:
    source, target, _ = resolve_local_folder_path(archive, layout, folder_id)
    if target.exists() and not target.is_dir():
        raise FileExistsError("本机目标位置已存在同名文件")
    created = 0 if target.exists() else 1
    target.mkdir(parents=True, exist_ok=True)
    return source, created


def _normalize_sync_relative(value: str, label: str) -> str:
    raw = str(value or "").replace("\\", "/").strip()
    parts = raw.split("/")
    if raw.startswith("/") or re.match(r"^[A-Za-z]:/", raw) or ".." in parts:
        raise ValueError(f"{label}路径不能指向同步目录之外")
    return normalize_archive_path(raw)


def sync_uploads_to_local(
    archive: dict,
    layout: dict,
    folder_id: str,
    root_name: str,
    directory_paths: list[str],
    uploads: list[tuple[str, bytes]],
) -> tuple[dict, int, int]:
    source, base_target, _ = resolve_local_folder_path(archive, layout, folder_id)
    target = base_target
    if root_name.strip():
        target = (base_target / _folder_name(root_name)).resolve()
    root = Path(source["root_path"]).expanduser().resolve()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise ValueError("同步目标超出本地来源范围") from exc

    normalized_directories = [_normalize_sync_relative(item, "上传目录") for item in directory_paths]
    normalized_directories = [item for item in normalized_directories if item]
    normalized_uploads = [(_normalize_sync_relative(path, "上传文件"), data) for path, data in uploads]
    normalized_uploads = [(path, data) for path, data in normalized_uploads if path]
    planned_files: list[tuple[Path, bytes]] = []
    seen: set[str] = set()
    for relative, data in normalized_uploads:
        destination = (target / relative).resolve()
        try:
            destination.relative_to(target)
        except ValueError as exc:
            raise ValueError("上传文件路径超出同步目录") from exc
        key = str(destination).casefold()
        if key in seen:
            raise FileExistsError(f"上传内容中存在同名文件：{Path(relative).name}")
        seen.add(key)
        if destination.exists():
            raise FileExistsError(f"本机目标目录已存在同名文件：{Path(relative).name}")
        planned_files.append((destination, data))

    planned_directories = {target}
    for relative in normalized_directories:
        destination = (target / relative).resolve()
        try:
            destination.relative_to(target)
        except ValueError as exc:
            raise ValueError("上传目录路径超出同步目录") from exc
        if destination.exists() and not destination.is_dir():
            raise FileExistsError(f"本机目标位置已存在同名文件：{Path(relative).name}")
        planned_directories.add(destination)
    planned_directories.update(destination.parent for destination, _ in planned_files)
    for directory in list(planned_directories):
        current = directory
        while current != root:
            planned_directories.add(current)
            current = current.parent
    ordered_directories = sorted(planned_directories, key=lambda item: len(item.parts))
    created_directories: list[Path] = []
    created_files: list[Path] = []
    try:
        for directory in ordered_directories:
            if not directory.exists():
                directory.mkdir()
                created_directories.append(directory)
        for destination, data in planned_files:
            destination.write_bytes(data)
            created_files.append(destination)
    except OSError:
        for destination in reversed(created_files):
            try:
                destination.unlink(missing_ok=True)
            except OSError:
                pass
        for directory in reversed(created_directories):
            try:
                directory.rmdir()
            except OSError:
                pass
        raise
    return source, len(created_files), len(created_directories)


def create_composition(payload: dict) -> dict:
    timestamp = utc_now()
    return {**payload, "id": str(uuid4()), "version": 1, "created_at": timestamp, "updated_at": timestamp}


def update_composition(record: dict, payload: dict) -> dict:
    if all(record.get(key) == value for key, value in payload.items()):
        return record
    record.update({**payload, "version": record.get("version", 1) + 1, "updated_at": utc_now()})
    return record


def composition_summary(record: dict) -> dict:
    return {
        "id": record["id"], "title": record["title"], "archive_id": record.get("archive_id"),
        "unit_id": record.get("unit_id"), "version": record.get("version", 1),
        "block_count": len(record.get("blocks", [])), "updated_at": record["updated_at"],
    }


def composition_markdown(record: dict) -> str:
    parts = [f"# {record['title']}"]
    for block in record.get("blocks", []):
        source = f"\n\n> 来源：{block.get('source_name')} · {block.get('locator')}" if block.get("source_name") else ""
        parts.append(f"## {block['title']}\n\n{block.get('content', '')}{source}")
    return "\n\n".join(parts).strip() + "\n"


def composition_html(record: dict) -> str:
    sections = []
    for block in record.get("blocks", []):
        content = html.escape(block.get("content", "")).replace("\n", "<br>")
        source = html.escape(" · ".join(filter(None, (block.get("source_name"), block.get("locator")))))
        sections.append(f"<section><h2>{html.escape(block['title'])}</h2><div>{content}</div>{f'<small>来源：{source}</small>' if source else ''}</section>")
    return f"<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'><title>{html.escape(record['title'])}</title><style>body{{max-width:900px;margin:40px auto;padding:0 28px;color:#172033;font:16px/1.75 Arial,'Microsoft YaHei',sans-serif}}h1{{font-size:28px;border-bottom:2px solid #246bce;padding-bottom:14px}}h2{{font-size:19px;margin-top:30px}}section{{border-bottom:1px solid #dde4ee;padding-bottom:22px}}small{{display:block;margin-top:14px;color:#667085}}</style></head><body><h1>{html.escape(record['title'])}</h1>{''.join(sections)}</body></html>"


def composition_docx(record: dict) -> bytes:
    document = Document()
    title = document.add_heading(record["title"], level=0)
    title.alignment = 1
    for block in record.get("blocks", []):
        document.add_heading(block["title"], level=1)
        for line in (block.get("content") or "").splitlines() or [""]:
            paragraph = document.add_paragraph(line)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10.5)
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        if block.get("source_name"):
            paragraph = document.add_paragraph(f"来源：{block['source_name']} · {block.get('locator', '')}")
            paragraph.style = document.styles["Caption"]
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def import_composition(filename: str, data: bytes, document_store: Path) -> dict:
    extension = Path(filename).suffix.lower()
    if extension == ".json":
        parsed_json = json.loads(data.decode("utf-8-sig"))
        if isinstance(parsed_json, dict) and isinstance(parsed_json.get("blocks"), list):
            payload = {key: parsed_json.get(key) for key in ("title", "archive_id", "unit_id", "blocks")}
            return create_composition(payload)
    if extension not in {".docx", ".pdf", ".pptx", ".md", ".txt"}:
        raise ValueError("仅支持导入 DOCX、PDF、PPTX、Markdown、TXT 或中台 JSON")
    parsed = parse_document(filename, data)
    persist_original(document_store, parsed["document_id"], filename, data)
    block = {
        "id": str(uuid4()), "source_block_id": None, "kind": "imported", "title": filename,
        "content": parsed.get("raw_text", ""), "source_name": filename,
        "locator": f"document:{parsed['document_id']}",
    }
    record = create_composition({"title": Path(filename).stem, "archive_id": None, "unit_id": None, "blocks": [block]})
    record["import_document_id"] = parsed["document_id"]
    record["import_original_url"] = f"/api/documents/{parsed['document_id']}/original"
    record["import_preview_url"] = f"/api/documents/{parsed['document_id']}/preview" if extension in {".docx", ".pdf", ".pptx"} else None
    return record
