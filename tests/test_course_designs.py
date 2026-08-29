from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from backend.course_designs import router as course_design_router
from backend.course_designs.models import CourseDesignAssemblyApply, CourseDesignContent, CourseDesignCreate
from backend.course_designs.service import (
    apply_assembly,
    assembly_sources,
    build_docx,
    create_design,
    inspect_docx_template,
    reference_detail,
    restore_source_snapshot,
    sync_run,
    update_design,
    validate_run_context,
)
from backend.course_designs.storage import list_designs, load_design, save_design
from backend.documents.storage import original_path, persist_original


MATERIAL_UNIT_ID = "aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa"
KNOWLEDGE_OUTLINE_ID = "bbbbbbbb-bbbb-4bbb-8bbb-bbbbbbbbbbbb"


def archive_record() -> dict:
    return {
        "id": "11111111-1111-4111-8111-111111111111",
        "course_title": "机电传动控制",
        "materials": [
            {
                "id": "22222222-2222-4222-8222-222222222222",
                "path": "教材/第3章.docx",
                "name": "第3章.docx",
                "extension": ".docx",
                "category": "textbook",
                "chapter": "第3章",
                "sha256": "abc123",
                "parse_status": "parsed",
                "document_id": "33333333-3333-4333-8333-333333333333",
                "preview_available": True,
                "character_count": 2400,
                "excerpt": "传感器把被测量转换为可用信号。",
            },
            {
                "id": "44444444-4444-4444-8444-444444444444",
                "path": "进度/教学进度表.xlsx",
                "name": "教学进度表.xlsx",
                "extension": ".xlsx",
                "category": "schedule",
                "chapter": None,
                "sha256": "def456",
                "parse_status": "parsed",
                "document_id": "55555555-5555-4555-8555-555555555555",
                "preview_available": False,
                "character_count": 300,
                "excerpt": "第3讲 传感器技术",
            },
        ],
        "schedule": [
            {
                "id": "schedule-3",
                "label": "第3讲",
                "content": "传感器定义、组成与性能指标",
                "chapter": "第3章",
                "source_material_id": "44444444-4444-4444-8444-444444444444",
            }
        ],
        "_documents": {
            "22222222-2222-4222-8222-222222222222": {
                "raw_text": "传感器把被测量转换为可用信号。\n传感器由敏感元件、转换元件和转换电路组成。",
                "sections": [{"id": "s1", "title": "传感器的定义", "start_offset": 0, "end_offset": 18}],
                "knowledge_points": [
                    {"title": "传感器的定义", "is_key_point": True, "difficulty_level": "中等"},
                    {"title": "传感器的组成", "is_key_point": True, "difficulty_level": "中等"},
                    {"title": "传感器的性能指标", "is_key_point": False, "difficulty_level": "较难"},
                ],
            },
            "44444444-4444-4444-8444-444444444444": {"raw_text": "第3讲 传感器技术", "sections": [], "knowledge_points": []},
        },
    }


def create_record() -> dict:
    return create_design(
        archive_record(),
        CourseDesignCreate(
            archive_id="11111111-1111-4111-8111-111111111111",
            chapter="第3章",
            schedule_id="schedule-3",
            material_ids=[
                "22222222-2222-4222-8222-222222222222",
                "44444444-4444-4444-8444-444444444444",
            ],
            primary_material_id="22222222-2222-4222-8222-222222222222",
        ),
    )


def knowledge_outline(version: int = 2) -> dict:
    return {
        "id": KNOWLEDGE_OUTLINE_ID,
        "version": version,
        "status": "confirmed",
        "title": "传感器基础知识大纲",
        "session": "第3次课 · 传感器定义、组成与性能指标",
        "knowledge_nodes": [
            {"title": "传感器的定义", "level": 1, "is_key_point": True, "is_difficult_point": False},
            {"title": "传感器的组成", "level": 2, "is_key_point": True, "is_difficult_point": False},
            {"title": "传感器性能指标", "level": 2, "is_key_point": False, "is_difficult_point": True},
        ],
        "source_references": [{"material_id": "source-1", "locator": "section:2.1"}],
        "scope_selection": {"teaching_item_ids": ["schedule-3"]},
    }


def outline_payload(version: int | None = 2) -> CourseDesignCreate:
    return CourseDesignCreate(
        archive_id="11111111-1111-4111-8111-111111111111",
        material_ids=["22222222-2222-4222-8222-222222222222"],
        primary_material_id="22222222-2222-4222-8222-222222222222",
        material_unit_id=MATERIAL_UNIT_ID,
        knowledge_outline_id=KNOWLEDGE_OUTLINE_ID,
        knowledge_outline_version=version,
    )


def test_knowledge_outline_drives_course_design_fields_and_lineage() -> None:
    record = create_design(archive_record(), outline_payload(), knowledge_outline())

    assert record["content"]["topic"] == "传感器基础知识大纲"
    assert record["title"] == "传感器基础知识大纲"
    assert record["content"]["session_label"].startswith("第3次课")
    assert record["content"]["knowledge_points"] == ["传感器的定义", "传感器的组成", "传感器性能指标"]
    assert record["content"]["key_points"] == ["传感器的定义", "传感器的组成"]
    assert record["content"]["difficult_points"] == ["传感器性能指标"]
    assert record["material_unit_id"] == MATERIAL_UNIT_ID
    assert record["knowledge_outline_version"] == 2
    structured = next(
        item for item in record["source_references"]
        if item["locator"].startswith("material-unit:")
    )
    assert structured["layer"] == "structured"
    assert structured["locator"] == (
        f"material-unit:{MATERIAL_UNIT_ID}:knowledge-outline:{KNOWLEDGE_OUTLINE_ID}:v2"
    )


def test_knowledge_outline_does_not_invent_unmarked_difficult_points() -> None:
    outline = knowledge_outline()
    for node in outline["knowledge_nodes"]:
        node["is_difficult_point"] = False

    record = create_design(archive_record(), outline_payload(), outline)

    assert record["content"]["difficult_points"] == []


def test_current_material_unit_nodes_are_mapped_into_course_design() -> None:
    outline = knowledge_outline()
    outline["nodes"] = outline.pop("knowledge_nodes")
    outline.pop("session")
    outline["selected_session_ids"] = [f"schedule:{MATERIAL_UNIT_ID}:schedule-3"]

    resolved = course_design_router._resolve_knowledge_outline(
        outline_payload(),
        {
            "id": MATERIAL_UNIT_ID,
            "archive_id": archive_record()["id"],
            "knowledge_outlines": [outline],
        },
        archive_record(),
    )
    record = create_design(archive_record(), outline_payload(), resolved)

    assert record["content"]["session_label"] == "传感器定义、组成与性能指标"
    assert record["content"]["knowledge_points"] == ["传感器的定义", "传感器的组成", "传感器性能指标"]
    assert record["content"]["key_points"] == ["传感器的定义", "传感器的组成"]
    assert record["source_references"][-1]["character_count"] > 0


def test_create_design_without_outline_keeps_legacy_material_behavior() -> None:
    record = create_record()

    assert record["material_unit_id"] is None
    assert record["knowledge_outline_id"] is None
    assert record["content"]["topic"] == "传感器定义、组成与性能指标"
    assert not any(
        item["locator"].startswith("material-unit:")
        for item in record["source_references"]
    )


def test_create_course_design_api_resolves_latest_outline_version(monkeypatch, tmp_path: Path) -> None:
    unit = {
        "id": MATERIAL_UNIT_ID,
        "archive_id": archive_record()["id"],
        "knowledge_outlines": [knowledge_outline(1), knowledge_outline(2)],
    }
    monkeypatch.setattr(course_design_router, "load_archive", lambda _root, _id: archive_record())
    monkeypatch.setattr(course_design_router, "load_material_unit", lambda _root, _id: unit)
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            material_unit_store_path=tmp_path / "units",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).post(
        "/api/course-designs",
        json=outline_payload(version=None).model_dump(exclude_none=True),
    )

    assert response.status_code == 201
    assert response.json()["knowledge_outline_version"] == 2
    assert response.json()["content"]["topic"] == "传感器基础知识大纲"


def test_create_course_design_api_rejects_missing_outline_version(monkeypatch, tmp_path: Path) -> None:
    unit = {
        "id": MATERIAL_UNIT_ID,
        "archive_id": archive_record()["id"],
        "knowledge_outlines": [knowledge_outline(1)],
    }
    monkeypatch.setattr(course_design_router, "load_archive", lambda _root, _id: archive_record())
    monkeypatch.setattr(course_design_router, "load_material_unit", lambda _root, _id: unit)
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            material_unit_store_path=tmp_path / "units",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).post(
        "/api/course-designs",
        json=outline_payload(version=9).model_dump(exclude_none=True),
    )

    assert response.status_code == 404
    assert "v9" in response.json()["detail"]


def test_design_preserves_original_extracted_and_structured_lineage() -> None:
    record = create_record()
    layers = [item["layer"] for item in record["source_references"]]
    assert layers.count("original") == 2
    assert layers.count("extracted") == 2
    assert layers.count("structured") == 1
    original = next(item for item in record["source_references"] if item["layer"] == "original")
    assert original["sha256"] == "abc123"
    assert original["original_url"].endswith("/original")
    assert record["content"]["topic"] == "传感器定义、组成与性能指标"


def test_extracted_reference_returns_full_text_and_sections() -> None:
    record = create_record()
    extracted = next(
        item for item in record["source_references"]
        if item["layer"] == "extracted" and item["material_id"] == record["primary_material_id"]
    )
    detail = reference_detail(record, archive_record(), extracted["id"])
    assert "敏感元件" in detail["content"]
    assert detail["sections"][0]["title"] == "传感器的定义"


def test_update_creates_version_snapshot() -> None:
    record = create_record()
    content = CourseDesignContent.model_validate(record["content"])
    content.topic = "传感器技术教学设计"
    updated = update_design(record, content, "reviewed", None)
    assert updated["version"] == 2
    assert updated["status"] == "reviewed"
    assert [item["version"] for item in updated["_versions"]] == [1, 2]


def test_sync_run_adds_generated_reference_and_framework_fields() -> None:
    record = create_record()
    run = {
        "id": "66666666-6666-4666-8666-666666666666",
        "objective": "第3章传感器技术",
        "status": "completed",
        "final_output": "教学成果",
        "teaching_data": {
            "archive_id": record["archive_id"],
            "design_id": record["id"],
            "document_id": "33333333-3333-4333-8333-333333333333",
            "content_analysis": {"key_points": ["定义", "组成"], "difficult_points": ["性能指标"]},
            "teaching_framework": {
                "learning_objectives": ["解释传感器定义"],
                "strategies": ["案例对比"],
                "stages": [{"name": "案例导入", "purpose": "建立问题", "activity": "比较两类传感器"}],
                "assessment": ["完成方案比较表"],
            },
        },
    }
    updated = sync_run(record, run)
    assert updated["run_id"] == run["id"]
    assert updated["content"]["objectives"] == ["解释传感器定义"]
    assert "案例导入" in updated["content"]["teaching_process"]
    assert any(item["layer"] == "generated" for item in updated["source_references"])


@pytest.mark.parametrize(
    "teaching_data",
    [
        {"design_id": "77777777-7777-4777-8777-777777777777"},
        {"archive_id": "88888888-8888-4888-8888-888888888888"},
        {"document_id": "99999999-9999-4999-8999-999999999999"},
        {"design_id": "77777777-7777-4777-8777-777777777777", "archive_id": "88888888-8888-4888-8888-888888888888"},
        {"design_id": "77777777-7777-4777-8777-777777777777", "document_id": "99999999-9999-4999-8999-999999999999"},
        {"archive_id": "88888888-8888-4888-8888-888888888888", "document_id": "99999999-9999-4999-8999-999999999999"},
        {"design_id": "77777777-7777-4777-8777-777777777777", "archive_id": "88888888-8888-4888-8888-888888888888", "document_id": "99999999-9999-4999-8999-999999999999"},
        {},
    ],
)
def test_run_context_rejects_eight_mismatch_shapes(teaching_data: dict) -> None:
    with pytest.raises(ValueError):
        validate_run_context(create_record(), {"teaching_data": teaching_data})


def test_legacy_run_context_accepts_matching_source_document() -> None:
    validate_run_context(
        create_record(),
        {"teaching_data": {"document_id": "33333333-3333-4333-8333-333333333333"}},
    )


def test_standard_docx_is_editable_and_contains_sources() -> None:
    record = create_record()
    data, mode = build_docx(record)
    document = Document(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert mode == "standard-template"
    assert "课程教学设计" in text
    assert "数据来源" in text
    assert len(data) > 20_000


def test_source_docx_template_is_filled(tmp_path: Path) -> None:
    template = Document()
    table = template.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = "待填写"
    table.cell(1, 0).text = "教学重点"
    table.cell(1, 1).text = "待填写"
    template_path = tmp_path / "template.docx"
    template.save(template_path)
    data, mode = build_docx(create_record(), template_path)
    result = Document(BytesIO(data))
    assert mode == "source-template"
    assert "准确说明" in result.tables[0].cell(0, 1).text
    assert "传感器的定义" in result.tables[0].cell(1, 1).text


def test_source_template_preserves_run_style_and_replaces_split_token(tmp_path: Path) -> None:
    template = Document()
    paragraph = template.add_paragraph()
    first = paragraph.add_run("{{teach")
    first.bold = True
    first.font.name = "Arial"
    paragraph.add_run("ing_process}}")
    template_path = tmp_path / "split-token.docx"
    template.save(template_path)

    report = inspect_docx_template(template_path, CourseDesignContent.model_validate(create_record()["content"]))
    data, mode = build_docx(create_record(), template_path)
    rendered = Document(BytesIO(data))

    assert report["compatible"] is True
    assert "teaching_process" in report["matched_fields"]
    assert mode == "source-template"
    assert "核心概念和原理分析" in rendered.paragraphs[0].text
    assert rendered.paragraphs[0].runs[0].bold is True
    assert rendered.paragraphs[0].runs[0].font.name == "Arial"


def test_incompatible_template_requires_explicit_standard_fallback(tmp_path: Path) -> None:
    template = Document()
    template.add_paragraph("普通通知模板，没有教案填充位置")
    template_path = tmp_path / "incompatible.docx"
    template.save(template_path)

    with pytest.raises(ValueError, match="未识别到"):
        build_docx(create_record(), template_path)
    data, mode = build_docx(create_record(), template_path, preserve_source_format=False)

    assert mode == "standard-template"
    assert Document(BytesIO(data)).paragraphs


def test_template_labels_do_not_match_guidance_or_similar_words(tmp_path: Path) -> None:
    template = Document()
    table = template.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "授课时间"
    table.cell(0, 1).text = "2026年春季"
    table.cell(1, 0).text = "教学目的与要求"
    table.cell(1, 1).text = "待填"
    table.cell(2, 0).text = "建议使用形成性评价检查目标"
    table.cell(2, 1).text = "说明文字"
    template_path = tmp_path / "similar-labels.docx"
    template.save(template_path)

    data, mode = build_docx(create_record(), template_path)
    rendered = Document(BytesIO(data)).tables[0]

    assert mode == "source-template"
    assert rendered.cell(0, 1).text == "2026年春季"
    assert "准确说明" in rendered.cell(1, 1).text
    assert rendered.cell(2, 1).text == "说明文字"


def test_merged_template_cell_is_filled_once_without_blank_paragraphs(tmp_path: Path) -> None:
    template = Document()
    table = template.add_table(rows=2, cols=4)
    table.cell(0, 0).merge(table.cell(0, 2)).text = "教学过程"
    target = table.cell(1, 0).merge(table.cell(1, 3))
    target.text = "待填"
    template_path = tmp_path / "merged.docx"
    template.save(template_path)

    content = CourseDesignContent.model_validate(create_record()["content"])
    report = inspect_docx_template(template_path, content)
    data, mode = build_docx(create_record(), template_path)
    rendered_cell = Document(BytesIO(data)).tables[0].cell(1, 0)

    assert mode == "source-template"
    assert report["replacement_count"] == 1
    assert rendered_cell.text == content.teaching_process
    assert len(rendered_cell.paragraphs) == len(content.teaching_process.splitlines())


def test_export_docx_api_returns_downloadable_word(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    expected, _ = build_docx(record)
    monkeypatch.setattr(course_design_router, "load_design", lambda _store, _design_id: record)
    monkeypatch.setattr(
        course_design_router,
        "build_docx",
        lambda _record, _template=None, _preserve=True: (expected, "standard-template"),
    )
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            document_store_path=tmp_path / "documents",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).post(
        f"/api/course-designs/{record['id']}/export.docx",
        json={"filename": "sensor-course-design.docx"},
    )

    assert response.status_code == 200
    assert response.headers["x-template-mode"] == "standard-template"
    assert response.headers["x-export-id"]
    assert response.headers["x-document-id"]
    assert "filename*=UTF-8''" in response.headers["content-disposition"]
    assert Document(BytesIO(response.content)).paragraphs
    saved = load_design(tmp_path / "designs", record["id"])
    export = saved["exports"][0]
    assert export["design_version"] == record["version"]
    assert export["size"] == len(expected)
    assert original_path(tmp_path / "documents", export["document_id"]).read_bytes() == expected


def test_template_inspection_and_export_history_api(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    save_design(tmp_path / "designs", record)
    template = Document()
    table = template.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = "待填"
    template_bytes = BytesIO()
    template.save(template_bytes)
    template_document_id = "77777777-7777-4777-8777-777777777777"
    persist_original(tmp_path / "documents", template_document_id, "lesson.docx", template_bytes.getvalue())
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            document_store_path=tmp_path / "documents",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)
    client = TestClient(api)

    inspection = client.post(
        f"/api/course-designs/{record['id']}/template-inspection",
        json={"template_document_id": template_document_id},
    )
    exported = client.post(
        f"/api/course-designs/{record['id']}/export.docx",
        json={"template_document_id": template_document_id, "filename": "教案\u0000测试.docx"},
    )
    history = client.get(f"/api/course-designs/{record['id']}/exports")

    assert inspection.status_code == 200
    assert inspection.json()["compatible"] is True
    assert "objectives" in inspection.json()["matched_fields"]
    assert exported.status_code == 200
    assert history.status_code == 200
    assert history.json()["items"][0]["filename"] == "教案测试.docx"
    assert history.json()["items"][0]["sha256"]


def test_delete_design_removes_persisted_exports(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    document_id = "88888888-8888-4888-8888-888888888888"
    persist_original(tmp_path / "documents", document_id, "export.docx", b"word-bytes")
    record["exports"] = [{"document_id": document_id}]
    save_design(tmp_path / "designs", record)
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            document_store_path=tmp_path / "documents",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).delete(f"/api/course-designs/{record['id']}")

    assert response.status_code == 204
    assert not (tmp_path / "documents" / document_id).exists()


def test_delete_single_export_keeps_design_and_removes_document(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    document_id = "99999999-9999-4999-8999-999999999999"
    export_id = "aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa"
    persist_original(tmp_path / "documents", document_id, "export.docx", b"word-bytes")
    record["exports"] = [{"id": export_id, "document_id": document_id}]
    save_design(tmp_path / "designs", record)
    monkeypatch.setattr(
        course_design_router,
        "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            document_store_path=tmp_path / "documents",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).delete(f"/api/course-designs/{record['id']}/exports/{export_id}")

    assert response.status_code == 204
    assert load_design(tmp_path / "designs", record["id"])["exports"] == []
    assert not (tmp_path / "documents" / document_id).exists()


def test_design_storage_round_trip(tmp_path: Path) -> None:
    record = create_record()
    save_design(tmp_path, record)
    assert load_design(tmp_path, record["id"])["content"]["course_name"] == "机电传动控制"
    assert list_designs(tmp_path)[0]["id"] == record["id"]


def test_syllabus_requirements_are_snapshotted_and_fill_design_fields() -> None:
    outline = knowledge_outline()
    outline["requirements"] = [
        {"id": "objective-1", "category": "objective", "category_label": "课程目标", "title": "目标", "content": "能够解释传感器的组成与信号转换过程", "score": 0.9, "reason": "讲次匹配", "recommended": True, "evidence": {"source_type": "syllabus", "material_id": "source-2", "quote": "解释传感器组成", "locator": "section:goal"}},
        {"id": "key-1", "category": "key_point", "category_label": "重点", "title": "重点", "content": "传感器性能指标", "score": 0.8, "reason": "讲次匹配", "recommended": True, "evidence": {"source_type": "syllabus", "material_id": "source-2", "quote": "性能指标", "locator": "section:key"}},
    ]

    record = create_design(archive_record(), outline_payload(), outline)

    assert record["content"]["objectives"] == ["能够解释传感器的组成与信号转换过程"]
    assert record["content"]["key_points"][0] == "传感器性能指标"
    assert record["source_snapshot"]["syllabus_requirements"][0]["evidence"]["locator"] == "section:goal"
    assert record["source_snapshot"]["knowledge_nodes"][0]["title"] == "传感器的定义"


def test_sync_run_preserves_teacher_postscript_and_imports_ideological_elements() -> None:
    record = create_record()
    record["content"]["postscript"] = "教师课后填写的真实反思"
    run = {
        "id": "66666666-6666-4666-8666-666666666666", "objective": "传感器", "status": "completed",
        "teaching_data": {
            "design_id": record["id"],
            "teaching_framework": {"ideological_elements": [{"dimension": "科学精神", "content": "依据实验数据修正判断", "integration_method": "对比测量误差"}]},
        },
    }

    updated = sync_run(record, run, "# 教师审核稿")

    assert updated["content"]["postscript"] == "教师课后填写的真实反思"
    assert updated["content"]["ideological_elements"] == ["科学精神；依据实验数据修正判断；对比测量误差"]


def test_assembly_sources_and_apply_targeted_teacher_answer() -> None:
    record = create_record()
    record["source_snapshot"] = {
        "schedule": [{"id": "lesson-3", "label": "第3次课", "content": "传感器定义与性能指标"}],
        "syllabus_requirements": [{"id": "goal", "category": "objective", "category_label": "课程目标", "content": "能分析传感器性能", "evidence": {"label": "课程大纲", "locator": "section:goal"}}],
        "knowledge_nodes": [],
    }
    run = {
        "id": "run-1", "teaching_data": {"messages": [{
            "id": "answer-1", "agent_name": "课程教师", "agent_type": "teacher",
            "phase": "teacher_answer", "iteration": 2, "content": "通过灵敏度与线性度对比解释选型。",
        }]},
    }
    sources = assembly_sources(record, run)
    answer = next(item for item in sources if item["kind"] == "teacher_message")

    updated = apply_assembly(
        record,
        CourseDesignAssemblyApply(base_version=record["version"], source_ids=[answer["id"]], target_field="teaching_process", mode="append"),
        sources,
    )

    assert "通过灵敏度与线性度对比解释选型" in updated["content"]["teaching_process"]
    assert updated["content_insertions"][-1]["locator"].endswith("teacher_answer:answer-1")
    assert updated["content_insertions"][-1]["target_field"] == "teaching_process"


def test_restore_legacy_design_snapshot_uses_exact_outline_version() -> None:
    record = create_record()
    record.update({"material_unit_id": "unit-1", "knowledge_outline_id": "outline-1", "knowledge_outline_version": 1})
    record.pop("source_snapshot", None)
    unit = {"knowledge_outlines": [
        {"id": "outline-1", "version": 1, "selected_session_ids": ["schedule:lesson-3"], "requirements": [{"id": "old-goal"}], "nodes": [{"id": "node-old", "title": "旧版本"}]},
        {"id": "outline-1", "version": 2, "selected_session_ids": [], "requirements": [{"id": "new-goal"}], "nodes": [{"id": "node-new", "title": "新版本"}]},
    ]}

    restored = restore_source_snapshot(record, {"schedule": [{"id": "lesson-3", "label": "第3次课", "content": "传感器"}]}, unit)

    assert restored["source_snapshot"]["syllabus_requirements"] == [{"id": "old-goal"}]
    assert restored["source_snapshot"]["knowledge_nodes"][0]["title"] == "旧版本"


def test_shared_merged_target_combines_process_and_ideology_and_removes_fixed_height(tmp_path: Path) -> None:
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import Cm

    template = Document()
    table = template.add_table(rows=3, cols=4)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "教学过程"
    table.cell(0, 2).merge(table.cell(0, 3)).text = "课程思政元素"
    target = table.cell(1, 0).merge(table.cell(1, 3))
    target.text = "待填"
    table.rows[1].height = Cm(0.5)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    template_path = tmp_path / "shared-target.docx"
    template.save(template_path)
    record = create_record()
    record["content"]["ideological_elements"] = ["科学精神：以数据校核结论"]

    data, mode = build_docx(record, template_path)
    rendered = Document(BytesIO(data))
    value = rendered.tables[0].cell(1, 0).text

    assert mode == "source-template"
    assert "【教学过程】" in value and "核心概念和原理分析" in value
    assert "【课程思政】" in value and "科学精神" in value
    assert rendered.tables[0].rows[1].height_rule is None


def test_standard_template_keeps_all_upstream_and_teacher_sections() -> None:
    record = create_record()
    record["content"].update({
        "session_label": "第3次课：传感器定义与性能指标",
        "knowledge_points": ["传感器定义", "静态性能指标"],
        "ideological_elements": ["以实验数据校核结论，培养严谨求实意识"],
        "postscript": "学生对灵敏度与线性度的区别仍需巩固。",
    })

    data, mode = build_docx(record)
    rendered = Document(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    metadata = "\n".join(cell.text for table in rendered.tables for row in table.rows for cell in row.cells)

    assert mode == "standard-template"
    assert "第3次课：传感器定义与性能指标" in metadata
    assert "二、知识点" in text and "静态性能指标" in text
    assert "六、课程思政融入" in text and "严谨求实意识" in text
    assert "七、教学过程" in text and record["content"]["teaching_process"] in text
    assert "九、教学后记" in text and "学生对灵敏度与线性度的区别仍需巩固" in text


def test_sync_run_api_rejects_unfinished_session(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    save_design(tmp_path / "designs", record)

    class WorkflowService:
        async def get_run(self, _run_id):
            return SimpleNamespace(status="running")

    monkeypatch.setattr(
        course_design_router, "get_settings",
        lambda: SimpleNamespace(course_design_store_path=tmp_path / "designs"),
    )
    api = FastAPI()
    api.state.workflow_service = WorkflowService()
    api.include_router(course_design_router.router)

    response = TestClient(api).post(f"/api/course-designs/{record['id']}/sync-run/run-1")

    assert response.status_code == 409
    assert "尚未完成" in response.json()["detail"]


def test_assembly_sources_api_restores_legacy_outline(monkeypatch, tmp_path: Path) -> None:
    record = create_record()
    record.update({"material_unit_id": "unit-1", "knowledge_outline_id": "outline-1", "knowledge_outline_version": 1})
    record.pop("source_snapshot", None)
    save_design(tmp_path / "designs", record)
    unit = {"knowledge_outlines": [{
        "id": "outline-1", "version": 1, "selected_session_ids": ["schedule:lesson-3"],
        "requirements": [{"id": "goal", "category": "objective", "category_label": "课程目标", "content": "能够解释传感器原理", "evidence": {"label": "课程大纲", "locator": "section:goal"}}],
        "nodes": [{"id": "node-1", "title": "传感器原理", "level": 1}],
    }]}
    archive = archive_record()
    archive["schedule"] = [{"id": "lesson-3", "label": "第3次课", "content": "传感器原理", "source_material_id": "source-1"}]
    monkeypatch.setattr(course_design_router, "load_archive", lambda _root, _id: archive)
    monkeypatch.setattr(course_design_router, "load_material_unit", lambda _root, _id: unit)
    monkeypatch.setattr(
        course_design_router, "get_settings",
        lambda: SimpleNamespace(
            course_design_store_path=tmp_path / "designs",
            course_archive_store_path=tmp_path / "archives",
            material_unit_store_path=tmp_path / "units",
        ),
    )
    api = FastAPI()
    api.include_router(course_design_router.router)

    response = TestClient(api).get(f"/api/course-designs/{record['id']}/assembly-sources")

    assert response.status_code == 200
    kinds = {item["kind"] for item in response.json()["items"]}
    assert {"schedule", "syllabus", "knowledge_outline"} <= kinds
