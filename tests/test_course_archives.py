from io import BytesIO

import fitz
from docx import Document
from openpyxl import Workbook

from backend.course_archives.models import ArchiveManifestItem, PrepareArchiveRequest
from backend.course_archives.deletion import deletion_impact, public_deletion_impact
from backend.course_archives.service import (
    analyze_course_archive,
    append_course_archive_files,
    classify_material,
    extract_course_archive_materials,
    extract_chapter,
    extract_lesson,
    extract_version,
    prepare_archive_pack,
    remove_course_archive_materials,
)


def test_material_metadata_classification_and_markers():
    assert classify_material("课程/机电传动控制教学大纲.docx") == "syllabus"
    assert classify_material("第3章/教学记录_第4轮.md") == "teaching_record"
    assert classify_material("第7章/定时器仿真/index.html") == "interactive"
    assert classify_material("实验/电机控制.py") == "experiment"
    assert classify_material("课程/第4章电动机运行分析.pdf") == "textbook"
    assert extract_chapter("课程/第七章/课件.pptx") == "第7章"
    assert extract_lesson("第7章/第12讲_定时器.md") == "第12讲"
    assert extract_version("第2章/教学记录_第3轮_优化版.md") == "第3轮 / 优化版"


def test_archive_analysis_detects_exact_duplicates_and_teacher_habits(tmp_path):
    files = {
        "机电传动控制/教学大纲.md": b"# Mechanical transmission course\nOBE course objectives and 42 hours.",
        "机电传动控制/第2章/教材.md": "# 第2章 直流电动机\n## 2.1 工业电机案例\n工程参数与系统链对比。".encode(),
        "机电传动控制/第2章/教学记录_第1轮.md": "# 第2章教学记录\n本轮围绕同一知识点优化案例。".encode(),
        "机电传动控制/第2章/教学记录_第2轮.md": "# 第2章教学记录\n本轮围绕同一知识点优化案例。".encode(),
        "机电传动控制/第2章/内容审核报告.md": "# 审核报告\n检查公式、内容和版式。".encode(),
        "机电传动控制/第2章/实验/控制程序.py": "# 工程实验\nprint('motor')".encode(),
        "机电传动控制/第3章/第3章_OBE融合式教学设计.md": "# OBE 教学设计\n课程目标、任务和评价证据逐项对应。".encode(),
    }
    manifest = [ArchiveManifestItem(path=path, size=len(data)) for path, data in files.items()]
    record = analyze_course_archive("机电传动控制", manifest, list(files.items()), tmp_path / "documents")

    assert record["total_files"] == 7
    assert record["parsed_files"] == 7
    assert record["duplicate_groups"] == 1
    duplicate_records = [item for item in record["materials"] if item["duplicate_group"]]
    assert {item["name"] for item in duplicate_records} == {"教学记录_第1轮.md", "教学记录_第2轮.md"}
    assert record["chapters"][0]["label"] == "第2章"
    habit_keys = {habit["key"] for habit in record["habits"]}
    assert {"constraint_alignment", "same_scope_iteration", "engineering_cases", "obe_mapping", "quality_audit"} <= habit_keys
    assert "不得自动进入下一章节" in record["preparation_profile_prompt"]


def test_schedule_spreadsheet_is_parsed_into_session_entries(tmp_path):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "教学进度"
    sheet.append(["周次", "教学内容", "学时"])
    sheet.append(["第3周", "第2章 直流电动机机械特性", 2])
    stream = BytesIO()
    workbook.save(stream)
    data = stream.getvalue()
    path = "机电传动控制/教学进度表.xlsx"
    record = analyze_course_archive(
        "机电传动控制",
        [ArchiveManifestItem(path=path, size=len(data))],
        [(path, data)],
        tmp_path / "documents",
    )

    assert record["materials"][0]["parse_status"] == "parsed"
    assert any(entry["chapter"] == "第2章" and "直流电动机" in entry["content"] for entry in record["schedule"])
    structured = record["_documents"][record["materials"][0]["id"]]["structured_content"]
    assert structured["kind"] == "spreadsheet"
    assert structured["sheets"][0]["rows"][0] == ["周次", "教学内容", "学时"]
    assert structured["sheets"][0]["column_count"] == 3


def test_archive_pdf_index_skips_ocr_and_marks_scanned_pages(tmp_path):
    pdf = fitz.open()
    text_page = pdf.new_page()
    text_page.insert_text((72, 72), "Chapter 2 motor control native teaching text")
    pdf.new_page()
    data = pdf.tobytes()
    pdf.close()
    path = "课程/第2章/教材.pdf"
    record = analyze_course_archive(
        "机电传动控制",
        [ArchiveManifestItem(path=path, size=len(data))],
        [(path, data)],
        tmp_path / "documents",
    )
    material = record["materials"][0]
    parsed = record["_documents"][material["id"]]

    assert material["parse_status"] == "parsed"
    assert parsed["extraction_report"]["engine"] == "PyMuPDF 快速索引（无 OCR）"
    assert parsed["extraction_report"]["scanned_page_count"] == 1
    assert parsed["extraction_report"]["ocr_page_count"] == 0


def test_archive_ocr_docx_is_fast_indexed_as_textbook(tmp_path):
    document = Document()
    document.add_heading("第4章 继电接触器控制", level=1)
    document.add_paragraph("低压电器的结构、工作原理和工程选型。")
    document.add_paragraph("第四章 继电接触器控制... ...43")
    document.add_paragraph("一、接触器的结构与工作原理")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "器件"
    table.cell(0, 1).text = "用途"
    table.cell(1, 0).text = "接触器"
    table.cell(1, 1).text = "主回路通断"
    stream = BytesIO()
    document.save(stream)
    data = stream.getvalue()
    path = "课程/第4章/机电传动控制 第2版第4章(OCR).docx"
    record = analyze_course_archive(
        "机电传动控制",
        [ArchiveManifestItem(path=path, size=len(data))],
        [(path, data)],
        tmp_path / "documents",
    )
    material = record["materials"][0]
    parsed = record["_documents"][material["id"]]

    assert material["category"] == "textbook"
    assert material["parse_status"] == "parsed"
    assert parsed["extraction_report"]["engine"] == "python-docx 快速索引（无图片 OCR）"
    assert "接触器" in parsed["raw_text"]
    assert parsed["course_name"].startswith("第4章 继电接触器控制")
    assert all("..." not in point["title"] for point in parsed["knowledge_points"])
    assert parsed["structured_content"]["kind"] == "word"
    assert parsed["structured_content"]["paragraphs"][0]["style"].startswith("Heading")
    assert parsed["structured_content"]["tables"][0]["rows"][1] == ["接触器", "主回路通断"]


def test_basename_fallback_never_associates_one_upload_to_ambiguous_paths(tmp_path):
    data = "# 教学记录\n同名文件正文".encode()
    manifest = [
        ArchiveManifestItem(path="课程/第1章/教学记录.md", size=len(data)),
        ArchiveManifestItem(path="课程/第2章/教学记录.md", size=len(data)),
    ]
    record = analyze_course_archive("课程", manifest, [("教学记录.md", data)], tmp_path / "documents")

    assert record["parsed_files"] == 0
    assert all(item["parse_status"] == "metadata_only" for item in record["materials"])


def test_prepare_pack_reuses_primary_document_and_supporting_context(tmp_path):
    files = {
        "课程/第3章/教材.md": "# 第3章 传感器技术\n## 传感器静态特性\n灵敏度与线性度。".encode(),
        "课程/教学大纲.md": "# 课程大纲\nOBE 课程目标与评价证据。".encode(),
        "课程/第3章/教学设计.md": "# 教学设计\n用工业机器人传感器案例进行比较教学。".encode(),
    }
    manifest = [ArchiveManifestItem(path=path, size=len(data)) for path, data in files.items()]
    record = analyze_course_archive("传感器课程", manifest, list(files.items()), tmp_path / "documents")
    by_name = {item["name"]: item for item in record["materials"]}
    requested = PrepareArchiveRequest(
        chapter="第3章",
        session_label="第5讲 传感器静态特性",
        material_ids=[item["id"] for item in record["materials"]],
        primary_material_id=by_name["教材.md"]["id"],
    )
    pack = prepare_archive_pack(record, requested)

    assert pack["primary_material_id"] == by_name["教材.md"]["id"]
    assert pack["parsed_document"]["file_name"] == "教材.md"
    assert "第5讲 传感器静态特性" in pack["context"]
    assert "教学设计.md" in pack["context"]
    assert len(pack["context"]) <= 7800
    assert pack["resources"][0]["role"] == "primary"


def test_upload_stores_original_and_extracts_only_when_requested(tmp_path):
    path = "课程/第1章/教材.md"
    data = "# 第一章 控制系统\n正文用于按需提取测试。".encode()
    record = analyze_course_archive(
        "控制课程",
        [ArchiveManifestItem(path=path, size=len(data))],
        [(path, data)],
        tmp_path / "documents",
        extract_uploads=False,
    )
    material = record["materials"][0]
    assert material["parse_status"] == "metadata_only"
    assert material["document_id"]
    assert record["parsed_files"] == 0
    assert record["_documents"] == {}

    extracted = extract_course_archive_materials(record, [material["id"]], tmp_path / "documents")
    assert extracted["materials"][0]["parse_status"] == "parsed"
    assert extracted["parsed_files"] == 1
    assert "按需提取测试" in extracted["_documents"][material["id"]]["raw_text"]


def test_requested_pdf_extraction_uses_full_document_parser(monkeypatch, tmp_path):
    path = "课程/第2章/扫描教材.pdf"
    data = b"%PDF-test-placeholder"
    record = analyze_course_archive(
        "控制课程",
        [ArchiveManifestItem(path=path, size=len(data))],
        [(path, data)],
        tmp_path / "documents",
        extract_uploads=False,
    )
    material = record["materials"][0]

    def fake_parse(filename, parsed_data):
        assert filename == "扫描教材.pdf"
        assert parsed_data == data
        return {
            "document_id": "replacement-id", "file_name": filename, "course_name": "控制课程",
            "raw_text": "第2章 动力学\n2.1 运动方程", "knowledge_points": [{"title": "2.1 运动方程"}],
            "sections": [{"id": "s1", "title": "2.1 运动方程", "level": 2, "preview": "转矩平衡"}],
            "extraction_report": {"engine": "PyMuPDF + RapidOCR v6", "quality_level": "high"},
            "character_count": 20, "processed_character_count": 20, "is_truncated": False,
        }

    monkeypatch.setattr("backend.course_archives.service.parse_document", fake_parse)
    extracted = extract_course_archive_materials(record, [material["id"]], tmp_path / "documents")

    assert extracted["materials"][0]["parse_status"] == "parsed"
    assert extracted["_documents"][material["id"]]["extraction_report"]["engine"] == "PyMuPDF + RapidOCR v6"
    assert extracted["_documents"][material["id"]]["document_id"] == material["document_id"]


def test_append_and_remove_platform_managed_originals(tmp_path):
    base_path = "课程/第1章/原有.txt"
    base_data = "原有资料正文".encode()
    record = analyze_course_archive(
        "控制课程",
        [ArchiveManifestItem(path=base_path, size=len(base_data))],
        [(base_path, base_data)],
        tmp_path / "documents",
        extract_uploads=False,
    )
    updated, added = append_course_archive_files(
        record,
        [("补充资料/案例.txt", "新增案例正文".encode())],
        tmp_path / "documents",
        "第1章",
    )
    assert updated["total_files"] == 2
    assert len(added) == 1
    assert added[0]["parse_status"] == "metadata_only"
    assert added[0]["chapter"] == "第1章"

    cleaned, document_ids = remove_course_archive_materials(updated, {added[0]["id"]})
    assert cleaned["total_files"] == 1
    assert document_ids == [added[0]["document_id"]]


def test_archive_deletion_impact_only_includes_explicitly_linked_records():
    archive_id = "11111111-1111-4111-8111-111111111111"
    design_id = "22222222-2222-4222-8222-222222222222"
    run_id = "33333333-3333-4333-8333-333333333333"
    archive = {
        "id": archive_id,
        "course_title": "临时测试课程",
        "materials": [
            {"document_id": "44444444-4444-4444-8444-444444444444"},
            {"document_id": "44444444-4444-4444-8444-444444444444"},
        ],
    }
    designs = [
        {"id": design_id, "archive_id": archive_id, "run_id": run_id, "exports": [{"document_id": "66666666-6666-4666-8666-666666666666"}]},
        {"id": "other-design", "archive_id": "other-archive"},
    ]
    runs = [
        {"id": run_id, "teaching_data": {}},
        {"id": "linked-by-data", "teaching_data": {"archive_id": archive_id}},
        {"id": "unrelated", "teaching_data": {"archive_id": "other-archive"}},
    ]
    compositions = [
        {"id": "composition", "archive_id": archive_id, "import_document_id": "55555555-5555-4555-8555-555555555555"},
        {"id": "other-composition", "archive_id": "other-archive"},
    ]
    layouts = [
        {"unit_id": f"{archive_id}:chapter"},
        {"unit_id": "other-archive:chapter"},
    ]

    impact = deletion_impact(archive, designs, runs, compositions, layouts)

    assert public_deletion_impact(impact) == {
        "archive_id": archive_id,
        "course_title": "临时测试课程",
        "material_count": 2,
        "document_count": 3,
        "design_count": 1,
        "composition_count": 1,
        "run_count": 2,
        "layout_count": 1,
    }
    assert {record["id"] for record in impact["_runs"]} == {run_id, "linked-by-data"}
